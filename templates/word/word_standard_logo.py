# templates/word/word_standard_logo.py
import os
import json
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT

# mapping English -> German for sex
SEX_MAP = {
    "MARE": "Stute",
    "GELDING": "Wallach", 
    "STALLION": "Hengst",
    "M": "Wallach",
    "F": "Stute",
    "": ""
}

WEEKDAY_MAP = {
    "Monday": "Montag", "Tuesday": "Dienstag", "Wednesday": "Mittwoch",
    "Thursday": "Donnerstag", "Friday": "Freitag", "Saturday": "Samstag",
    "Sunday": "Sonntag"
}

FIXED_WORD_TEMPLATE = "word_details_modern_simple.docx"

def get_country_name(ioc_code):
    """Gibt deutschen Ländernamen für IOC-Code zurück"""
    countries = {
        "GER": "Deutschland", "AUT": "Österreich", "SUI": "Schweiz",
        "NED": "Niederlande", "BEL": "Belgien", "FRA": "Frankreich",
        "GBR": "Großbritannien", "IRL": "Irland", "SWE": "Schweden",
        "NOR": "Norwegen", "DEN": "Dänemark", "FIN": "Finnland",
        "POL": "Polen", "CZE": "Tschechien", "HUN": "Ungarn",
        "ITA": "Italien", "ESP": "Spanien", "POR": "Portugal",
        "USA": "USA", "CAN": "Kanada", "BRA": "Brasilien",
        "AUS": "Australien", "NZL": "Neuseeland", "RSA": "Südafrika",
        "UAE": "Vereinigte Arabische Emirate", "QAT": "Katar",
        "BHR": "Bahrain", "KSA": "Saudi-Arabien", "MEX": "Mexiko",
        "ARG": "Argentinien", "CHI": "Chile", "COL": "Kolumbien",
        "JPN": "Japan", "CHN": "China", "KOR": "Südkorea",
    }
    return countries.get(str(ioc_code).upper(), str(ioc_code))

def set_light_gray_background(row):
    """Hellgraue Hintergrundfarbe (wie im PDF)"""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E8E8E8')
        tcPr.append(shd)


def _set_odd_even_headers(doc):
    """Setzt 'Gerade & ungerade Seiten unterschiedlich' in Word-Einstellungen"""
    settings = doc.settings.element
    # Prüfen ob bereits vorhanden
    if settings.find(qn('w:evenAndOddHeaders')) is None:
        even_odd = OxmlElement('w:evenAndOddHeaders')
        settings.insert(0, even_odd)


OUTPUT_DIR = "Ausgabe"

def _ensure_output_dir():
    """Stellt sicher, dass das Ausgabe-Verzeichnis existiert"""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

def _safe_get(d, key, default=""):
    if not d:
        return default
    return d.get(key, default) if isinstance(d, dict) else default

def _format_time(iso):
    """Formatiert ISO-Zeit zu HH:MM:SS"""
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", ""))
        return dt.strftime("%H:%M:%S")
    except Exception:
        # Fallback: erste 8 Zeichen für vollständige Zeit
        return str(iso)[:8] if len(str(iso)) >= 8 else str(iso)

def _format_header_datetime(iso):
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", ""))
        weekday = WEEKDAY_MAP.get(dt.strftime("%A"), dt.strftime("%A"))
        return f"{weekday}, {dt.strftime('%d.%m.%Y um %H:%M Uhr')}"
    except Exception:
        return str(iso)

def _process_information_text(text):
    """Verarbeitet informationText - entfernt führende \n und macht Text fett"""
    if not text:
        return ""
    text = text.lstrip('\n')
    return text

def _set_header_gray_background(row):
    """Set gray background for header row"""
    tr = row._tr
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D9D9D9')  # Light gray
        tcPr.append(shd)

def _set_table_borders(table):
    """Add borders to all table cells"""
    tbl = table._tbl
    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.tcPr
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell.append(tcPr)
            
            tcBorders = OxmlElement('w:tcBorders')
            
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)


def _add_sponsorship_footer(doc, show_sponsor_bar=True, show_banner=False):
    """Fügt Sponsorenleiste im Footer jeder Seite hinzu"""
    if not show_sponsor_bar:
        return
    sponsor_logo_path = os.path.join("logos", "sponsorenleiste.png")
    if os.path.exists(sponsor_logo_path):
        try:
            for section in doc.sections:
                footer = section.footer
                footer_para = footer.paragraphs[0]
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                footer_para.clear()
                footer_run = footer_para.add_run()
                footer_run.add_picture(sponsor_logo_path, width=Inches(7.5))
                footer_para.paragraph_format.space_after = Pt(0)
                footer_para.paragraph_format.space_before = Pt(0)
                # Bei Banner auch first_page_footer befüllen
                if show_banner and section.different_first_page_header_footer:
                    try:
                        _fpf = section.first_page_footer
                        _fpfp = _fpf.paragraphs[0] if _fpf.paragraphs else _fpf.add_paragraph()
                        _fpfp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        _fpfp.clear()
                        _fpfr = _fpfp.add_run()
                        _fpfr.add_picture(sponsor_logo_path, width=Inches(7.5))
                        _fpfp.paragraph_format.space_after = Pt(0)
                        _fpfp.paragraph_format.space_before = Pt(0)
                    except: pass
            
            print(f"WORD DRE 3 LOGO DEBUG: Sponsorenleiste im Footer hinzugefügt: {sponsor_logo_path}")
        except Exception as e:
            print(f"WORD DRE 3 LOGO DEBUG: Sponsorenleiste Footer-Fehler: {e}")
    else:
        print(f"WORD DRE 3 LOGO DEBUG: Sponsorenleiste nicht gefunden: {sponsor_logo_path}")



def _set_row_height_auto(row):
    """Set row height to auto (compact)"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is not None:
        trPr.remove(trHeight)
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '0')
    trHeight.set(qn('w:hRule'), 'auto')
    trPr.append(trHeight)

def _set_row_keep_together(row):
    """Prevent table row from breaking across pages"""
    tr = row._tr
    trPr = tr.trPr
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    
    # Add cantSplit property to prevent row from breaking across pages
    cantSplit = trPr.find(qn('w:cantSplit'))
    if cantSplit is None:
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)
    else:
        cantSplit.set(qn('w:val'), 'true')

def _set_compact_cell_format(paragraph):
    """Set compact formatting for table cells"""
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(1)
    paragraph.paragraph_format.line_spacing = 1.0

def _set_header_row_repeat(table):
    """Enable header row to repeat on each new page"""
    tbl = table._tbl
    
    # Get the first row (header row)
    if len(tbl.tr_lst) > 0:
        header_row = tbl.tr_lst[0]
        
        # Set the tblHeader property to repeat the header row on new pages
        trPr = header_row.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            header_row.insert(0, trPr)
        
        # Add or update tblHeader element
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is None:
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), 'true')
            trPr.append(tblHeader)
        else:
            tblHeader.set(qn('w:val'), 'true')
        
        print("WORD DEBUG: Header row repeat enabled")

def _set_judges_table_widths(table):
    """Set fixed column widths for judges table"""
    tbl = table._tbl
    
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblPr.remove(tblStyle)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    tblGrid = OxmlElement('w:tblGrid')
    judges_column_widths_twips = [360, 7200]
    
    for width in judges_column_widths_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)
    
    if tblPr.getparent() is not None:
        tblPr.getparent().insert(tblPr.getparent().index(tblPr) + 1, tblGrid)
    else:
        tbl.insert(0, tblGrid)
    
    for row in tbl.tr_lst:
        for i, cell in enumerate(row.tc_lst):
            if i < len(judges_column_widths_twips):
                tcPr = cell.tcPr
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell.insert(0, tcPr)
                
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(judges_column_widths_twips[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def _set_kompakt_column_widths(table):
    """Set table to use exact column widths for 6-column kompakt layout (Start|Zeit|KoNr|Reiter|Pferd|Ergeb)"""
    tbl = table._tbl
    
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblPr.remove(tblStyle)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    tblGrid = OxmlElement('w:tblGrid')
    # 6-column widths: Start | Zeit | KoNr. (breiter) | Reiter (breiter) | Pferd (breiter) | Ergeb. (breiter) - NUTZT GANZE SEITE
    kompakt_column_widths_twips = [600, 600, 900, 3600, 3600, 1500]  # KoNr von 600 auf 900 erweitert
    
    for width in kompakt_column_widths_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)
    
    if tblPr.getparent() is not None:
        tblPr.getparent().insert(tblPr.getparent().index(tblPr) + 1, tblGrid)
    else:
        tbl.insert(0, tblGrid)
    
    for row in tbl.tr_lst:
        for i, cell in enumerate(row.tc_lst):
            if i < len(kompakt_column_widths_twips):
                tcPr = cell.tcPr
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell.insert(0, tcPr)
                
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(kompakt_column_widths_twips[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def _get_judge_data_for_display_kompakt(judges):
    """
    Prozessiert Richter für die Anzeige - alle Positionen für die Richterliste in korrekter Reihenfolge
    """
    pos_map = {
        0: "E", 1: "H", 2: "C", 3: "M", 4: "B", 5: "K", 6: "V", 
        7: "S", 8: "R", 9: "P", 10: "F", 11: "A",
        "WARM_UP_AREA": "Aufsicht", "WATER_JUMP": "Wasser"
    }
    
    dressage_positions = ["E", "H", "C", "M", "B"]
    judges_by_position = {}  # Dictionary of Lists!
    other_judges = []
    
    for judge in judges:
        position = judge.get("position", "")
        if isinstance(position, int):
            pos_label = pos_map.get(position, str(position))
        else:
            pos_label = pos_map.get(str(position), str(position))
        
        judge_copy = judge.copy()
        judge_copy["pos_label"] = pos_label
        
        if pos_label in dressage_positions:
            # Füge zu Liste hinzu statt zu überschreiben
            if pos_label not in judges_by_position:
                judges_by_position[pos_label] = []
            judges_by_position[pos_label].append(judge_copy)
        else:
            other_judges.append(judge_copy)
    
    # Return judges in E H C M B order first, then others
    result = []
    for pos in dressage_positions:
        if pos in judges_by_position:
            # Füge ALLE Richter dieser Position hinzu
            result.extend(judges_by_position[pos])
    
    # Add other judges sorted by position label
    other_judges.sort(key=lambda j: j["pos_label"])
    result.extend(other_judges)
    
    return result
    
    
def render(starterlist: dict, filename: str):
    """
    Creates a Word document with 6-column kompakt layout (Start|Zeit|KoNr|Reiter|Pferd|Ergeb) - Standard version WITH LOGO
    """
    print("WORD STANDARD LOGO DEBUG: Starting render function")
    
    if starterlist is None:
        raise ValueError("starterlist is None")

    if isinstance(starterlist, str):
        try:
            starterlist = json.loads(starterlist)
        except Exception as e:
            raise ValueError(f"starterlist is a string but not valid JSON: {e}")

    _ensure_output_dir()
    # Druckoptionen auslesen
    print_options     = starterlist.get("printOptions", {})
    show_header       = print_options.get("show_header",      True)
    show_banner       = print_options.get("show_banner",      True)
    show_sponsor_bar  = print_options.get("show_sponsor_bar", True)
    show_title_opt    = print_options.get("show_title",       True)
    sponsor_top       = print_options.get("sponsor_top",      False)
    sponsor_bottom    = print_options.get("sponsor_bottom",   False)
    spacing_top_cm    = starterlist.get("spacingTopCm",   3.0)
    spacing_bottom_cm = starterlist.get("spacingBottomCm",2.0)
    
    # Use the fixed Word template as base
    template_path = os.path.join("templates", "word", FIXED_WORD_TEMPLATE)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Fixed Word template not found: {template_path}")
    
    # Output path
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    # Copy template to output
    shutil.copy2(template_path, output_path)
    
    # Open the copied template document
    print("WORD STANDARD LOGO DEBUG: Loading document from template")
    doc = Document(output_path)

    # Seitenformat + Banner
    _section = doc.sections[0]
    _section.page_height = Inches(11.69)
    _section.page_width  = Inches(8.27)
    _section.left_margin  = Inches(0.5)
    _section.right_margin = Inches(0.5)
    _MIN_TOP    = Inches(10 / 25.4)
    _MIN_BOTTOM = Inches(0.26)
    _SP_TOP     = Inches(spacing_top_cm / 2.54)
    _SP_BOTTOM  = Inches(spacing_bottom_cm / 2.54)

    if sponsor_top or sponsor_bottom:
        _section.different_odd_and_even_pages_header_footer = True
        _set_odd_even_headers(doc)
    if sponsor_top:
        _section.header_distance = Inches(0.2)
        _section.top_margin      = _SP_TOP
    else:
        _section.header_distance = Inches(0.2)
        _section.top_margin      = _MIN_TOP

    if sponsor_bottom:
        _section.footer_distance = _SP_BOTTOM
        _section.bottom_margin   = _SP_BOTTOM
    else:
        _section.footer_distance = Inches(0.24)
        _section.bottom_margin   = _MIN_BOTTOM
    if show_banner:
        _bp = "logos/banner.png"
        if os.path.exists(_bp):
            try:
                from PIL import Image as _BPIL
                _bi = _BPIL.open(_bp)
                _bw, _bh = _bi.size
                _bd = _bi.info.get('dpi', (96, 96))
                _bdx = float(_bd[0]) if isinstance(_bd, tuple) else float(_bd)
                if _bdx < 1: _bdx = 96.0
                _bh_in = _bh / _bdx
                _section.different_first_page_header_footer = True
                _section.header_distance = Inches(0)
                _section.top_margin = Inches(_bh_in + 0.04)
                _section.left_margin  = Inches(0.5)
                _section.right_margin = Inches(0.5)
                _fph = _section.first_page_header
                for _p in _fph.paragraphs: _p.clear()
                _hp = _fph.paragraphs[0]
                _hp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                _pPr = _hp._p.get_or_add_pPr()
                _sp = OxmlElement('w:spacing'); _sp.set(qn('w:before'),'0'); _sp.set(qn('w:after'),'220'); _pPr.append(_sp)  # 220 twips ≈ 3.9mm
                _ind = OxmlElement('w:ind'); _ind.set(qn('w:left'),str(-720)); _ind.set(qn('w:right'),'0'); _ind.set(qn('w:hanging'),'0'); _pPr.append(_ind)
                _hr = _hp.add_run()
                _hr.add_picture(_bp, width=Inches(8.27))
                _nh = _section.header; _nh.is_linked_to_previous = False
                for _p in _nh.paragraphs: _p.clear()
            except Exception as _be:
                print(f"Banner Fehler: {_be}")
    
    # Find and replace STARTER_TABLE placeholder
    print("WORD DEBUG: Looking for STARTER_TABLE placeholder")
    starter_table_found = False
    placeholder_paragraph = None
    
    for paragraph in doc.paragraphs:
        if "STARTER_TABLE" in paragraph.text:
            print("WORD DEBUG: Found STARTER_TABLE placeholder")
            placeholder_paragraph = paragraph
            starter_table_found = True
            break
    
    if not starter_table_found:
        print("WORD DEBUG: STARTER_TABLE placeholder not found, creating at end of document")
        placeholder_paragraph = doc.add_paragraph("STARTER_TABLE")

    # Extract data from starterlist for header generation
    show_title = starterlist.get("showTitle", "") if show_title_opt else ""
    comp_title = starterlist.get("competitionTitle", "")
    comp_number = starterlist.get("competitionNumber", "")
    comp_subtitle = starterlist.get("subtitle", "")
    comp_info = starterlist.get("informationText", "")
    location = starterlist.get("location", "")
    start_time = starterlist.get("start", "")
    division_number = starterlist.get("divisionNumber")

    # Handle division start times
    start_raw = None
    divisions = starterlist.get('divisions', [])
    
    if divisions and division_number is not None:
        try:
            div_num = int(division_number)
            for div in divisions:
                if div.get("number") == div_num:
                    division_start = div.get("start")
                    if division_start:
                        start_raw = division_start
                        break
        except (ValueError, TypeError):
            pass
    elif divisions and len(divisions) > 0:
        first_division = divisions[0]
        division_start = first_division.get("start")
        if division_start:
            start_raw = division_start
    
    if not start_raw:
        start_raw = start_time
    
    if start_raw:
        start_time = start_raw

    # Remove the placeholder paragraph first
    placeholder_paragraph._element.getparent().remove(placeholder_paragraph._element)

    if show_header:
        logo_path = starterlist.get("logoPath")
        logo_max_width_cm = starterlist.get("logoMaxWidthCm", 5.0)
        logo_inches = logo_max_width_cm / 2.54
        page_twips = int(7.27 * 1440)
        logo_twips = int((logo_inches + 0.1) * 1440)
        text_twips = page_twips - logo_twips

        if logo_path and os.path.exists(logo_path):
            h_table = doc.add_table(rows=1, cols=2)
            h_tbl = h_table._tbl
            h_tbl_pr = h_tbl.find(qn("w:tblPr"))
            if h_tbl_pr is None:
                h_tbl_pr = OxmlElement("w:tblPr"); h_tbl.insert(0, h_tbl_pr)
            h_borders = OxmlElement("w:tblBorders")
            for bn in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{bn}"); b.set(qn("w:val"), "none"); h_borders.append(b)
            h_tbl_pr.append(h_borders)
            h_tbl_w = OxmlElement("w:tblW")
            h_tbl_w.set(qn("w:w"), str(page_twips)); h_tbl_w.set(qn("w:type"), "dxa"); h_tbl_pr.append(h_tbl_w)
            tbl_grid = OxmlElement("w:tblGrid")
            for tw in [text_twips, logo_twips]:
                gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(tw)); tbl_grid.append(gc)
            h_tbl.insert(list(h_tbl).index(h_tbl_pr) + 1, tbl_grid)
            left_cell = h_table.rows[0].cells[0]
            right_cell = h_table.rows[0].cells[1]
            for cell, tw in zip([left_cell, right_cell], [text_twips, logo_twips]):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = OxmlElement("w:tcW"); tc_w.set(qn("w:w"), str(tw)); tc_w.set(qn("w:type"), "dxa"); tc_pr.append(tc_w)
            first_tracker = [True]
            def _add_h(cell, text, sz, bold, sa=4):
                if first_tracker[0]:
                    p = cell.paragraphs[0]; p.clear(); first_tracker[0] = False
                else:
                    p = cell.add_paragraph()
                r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(sa)
                return p
            if show_title: _add_h(left_cell, show_title, 14, True, sa=2)
            if comp_number or comp_title:
                _cl = f"Prüfung {comp_number}"
                if comp_title: _cl += f" \u2014 {comp_title}"
                if division_number:
                    try:
                        _dn = int(division_number)
                        if _dn > 0: _cl += f" - {_dn}. Abt."
                    except: pass
                _add_h(left_cell, _cl, 11, True, sa=4)
            if comp_info: _add_h(left_cell, comp_info.lstrip("\n"), 10, False, sa=4)
            if comp_subtitle: _add_h(left_cell, comp_subtitle, 10, False, sa=4)
            # Logo rechts
            r_para = right_cell.paragraphs[0]
            r_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            r_para.paragraph_format.space_before = Pt(0)
            r_para.paragraph_format.space_after = Pt(0)
            r_run = r_para.add_run()
            r_run.add_picture(logo_path, width=Inches(logo_inches))
            # behindDoc=1
            try:
                import copy as _copy
                _wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                _am = "http://schemas.openxmlformats.org/drawingml/2006/main"
                inline = r_run._r.find(f"{{{_wp}}}inline")
                if inline is not None:
                    anchor = OxmlElement("wp:anchor")
                    for k,v in [("distT","0"),("distB","0"),("distL","114300"),("distR","114300"),
                                 ("simplePos","0"),("relativeHeight","251658240"),("behindDoc","1"),
                                 ("locked","0"),("layoutInCell","1"),("allowOverlap","0")]:
                        anchor.set(k, v)
                    sp = OxmlElement("wp:simplePos"); sp.set("x","0"); sp.set("y","0"); anchor.append(sp)
                    posH = OxmlElement("wp:positionH"); posH.set("relativeFrom","margin")
                    alignH = OxmlElement("wp:align"); alignH.text = "right"; posH.append(alignH); anchor.append(posH)
                    posV = OxmlElement("wp:positionV"); posV.set("relativeFrom","margin")
                    alignV = OxmlElement("wp:align"); alignV.text = "top"; posV.append(alignV); anchor.append(posV)
                    extent = inline.find(f"{{{_wp}}}extent")
                    if extent is not None: anchor.append(_copy.deepcopy(extent))
                    eff = OxmlElement("wp:effectExtent")
                    for a in ["l","t","r","b"]: eff.set(a,"0")
                    anchor.append(eff)
                    anchor.append(OxmlElement("wp:wrapNone"))
                    docPr = inline.find(f"{{{_wp}}}docPr")
                    if docPr is not None: anchor.append(_copy.deepcopy(docPr))
                    anchor.append(OxmlElement("wp:cNvGraphicFramePr"))
                    a_graphic = inline.find(f"{{{_am}}}graphic")
                    if a_graphic is None:
                        for child in inline:
                            if child.tag.endswith("}graphic"): a_graphic = child; break
                    if a_graphic is not None: anchor.append(_copy.deepcopy(a_graphic))
                    drawing = inline.getparent()
                    if drawing is not None: drawing.replace(inline, anchor)
            except Exception as _be:
                print(f"behindDoc Fehler: {_be}")
        else:
            def _add_d(text, sz, bold, sa=4):
                p = doc.add_paragraph()
                r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(sa)
            if show_title: _add_d(show_title, 14, True, sa=2)
            if comp_number or comp_title:
                _cl = f"Prüfung {comp_number}"
                if comp_title: _cl += f" \u2014 {comp_title}"
                if division_number:
                    try:
                        _dn = int(division_number)
                        if _dn > 0: _cl += f" - {_dn}. Abt."
                    except: pass
                _add_d(_cl, 11, True, sa=4)
            if comp_info: _add_d(comp_info.lstrip("\n"), 10, False, sa=4)
            if comp_subtitle: _add_d(comp_subtitle, 10, False, sa=4)

    # Datum für Starterliste-Zeile
    _date_text = _format_header_datetime(start_time) if start_time else ""
    if _date_text and location:
        _date_text += f" - {location}"

    # "Starterliste" links + Datum rechtsbündig
    _PAGE_TWIPS = int(7.27 * 1440)
    p_sl = doc.add_paragraph()
    _pPr_sl = p_sl._p.get_or_add_pPr()
    # Paragraph-Einzug explizit auf 0
    _ind = OxmlElement('w:ind')
    _ind.set(qn('w:left'), '0')
    _ind.set(qn('w:right'), '0')
    _pPr_sl.append(_ind)
    # Tab-Stop rechts
    _tabs_sl = OxmlElement('w:tabs')
    _tab_sl  = OxmlElement('w:tab')
    _tab_sl.set(qn('w:val'), 'right')
    _tab_sl.set(qn('w:pos'), str(_PAGE_TWIPS))
    _tabs_sl.append(_tab_sl)
    _pPr_sl.append(_tabs_sl)
    p_sl.paragraph_format.space_before = Pt(2)
    p_sl.paragraph_format.space_after  = Pt(2)
    p_sl.paragraph_format.left_indent  = 0
    p_sl.paragraph_format.right_indent = 0
    run_sl = p_sl.add_run("Starterliste")
    run_sl.font.bold = True
    run_sl.font.size = Pt(12)
    if _date_text:
        p_sl.add_run("\t")
        run_sl_date = p_sl.add_run(_date_text)
        run_sl_date.font.bold = True
        run_sl_date.font.size = Pt(10)

    # Create table DIREKT nach kompaktem Header - 6 COLUMNS for kompakt layout (Start|Zeit|KoNr|Reiter|Pferd|Ergeb)
    print("WORD DEBUG: Creating table after compact header - 6 columns kompakt")
    table = doc.add_table(rows=0, cols=6)  # 6 columns: Start, Zeit, KoNr, Reiter, Pferd, Ergeb.
    
    # Try to set table style, but continue without if not available
    try:
        table.style = 'Table Grid'
        print("WORD DEBUG: Applied Table Grid style")
    except Exception as e:
        print(f"WORD DEBUG: Table Grid style not available: {e}, continuing without style")
    
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Add header row
    hdr_row = table.add_row()
    hdr_cells = hdr_row.cells
    
    headers = ["Start", "Zeit", "KoNr.", "Reiter", "Pferd", "Ergeb."]
    
    for i, header_text in enumerate(headers):
        if i >= len(hdr_cells):
            break
            
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        
        run = p.add_run(str(header_text))
        run.font.bold = True
        run.font.size = Pt(9)
        if i == 3 or i == 4:  # Reiter and Pferd columns - linksbündig
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Apply gray background to header row
    _set_header_gray_background(hdr_row)

    _set_kompakt_column_widths(table)
    
    # Enable header row repeat on new pages
    _set_header_row_repeat(table)

    # Process starters and breaks mit Abteilungslogik
    starters = starterlist.get("starters", [])
    breaks = starterlist.get("breaks", [])
    
    print(f"WORD DEBUG: Found {len(starters)} starters in starterlist")
    print(f"WORD DEBUG: Found {len(breaks)} breaks in starterlist")
    
    breaks_map = {}
    for br in breaks:
        try:
            after_num = br.get("afterNumberInCompetition")
            if after_num is None:
                k = -1
            else:
                k = int(after_num)
            breaks_map[k] = br
        except:
            continue

    # Gruppierung nach groupNumber für Abteilungen
    starters_by_group = {}
    for starter in starters:
        group_num = starter.get("groupNumber", 0)
        if group_num not in starters_by_group:
            starters_by_group[group_num] = []
        starters_by_group[group_num].append(starter)
    
    # Pause VOR erstem Starter (afterNumberInCompetition=0)
    row_counter = 0
    if 0 in breaks_map:
        break_info = breaks_map[0]
        total_seconds = int(break_info.get("totalSeconds", 0) or 0)
        info_text = break_info.get("informationText", "")
        if total_seconds > 0:
            hours = total_seconds // 3600
            mins = (total_seconds % 3600) // 60
            if hours > 0:
                break_text = f"Pause ({hours} h {mins:02d} min)"
            else:
                break_text = f"Pause ({mins} min)"
            if info_text:
                break_text = f"{break_text} - {info_text}"
        else:
            break_text = info_text or "Pause"
        br_row = table.add_row()
        _set_row_height_auto(br_row)
        _set_row_keep_together(br_row)
        p_br = br_row.cells[0].paragraphs[0]
        p_br.clear()
        run_br = p_br.add_run(break_text)
        run_br.font.size = Pt(10)
        run_br.font.bold = True
        p_br.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_br.paragraph_format.space_before = Pt(5)
        p_br.paragraph_format.space_after  = Pt(5)
        _set_compact_cell_format(p_br)
        for i in range(1, 6):
            br_row.cells[i].paragraphs[0].clear()
        cell_0 = br_row.cells[0]
        for i in range(1, 6):
            cell_0.merge(br_row.cells[i])
        tc_pr_b0 = cell_0._tc.get_or_add_tcPr()
        v_align0 = OxmlElement('w:vAlign')
        v_align0.set(qn('w:val'), 'center')
        tc_pr_b0.append(v_align0)
        tc_pr_b0 = cell_0._tc.get_or_add_tcPr()
        v_align0 = OxmlElement('w:vAlign')
        v_align0.set(qn('w:val'), 'center')
        tc_pr_b0.append(v_align0)
        # Zebra für Pause-0
        if row_counter % 2 == 1:
            set_light_gray_background(br_row)
        row_counter += 1

    # Gruppen sortiert verarbeiten
    for group_num in sorted(starters_by_group.keys()):
        group_starters = starters_by_group[group_num]
        
        # Abteilungsheader hinzufügen (wenn mehr als eine Gruppe)
        if len(starters_by_group) > 1 and group_num > 0:
            header_row = table.add_row()
            header_cells = header_row.cells
            
            # Abteilungstext linksbündig
            p_header = header_cells[0].paragraphs[0]
            p_header.clear()
            run = p_header.add_run(f"{group_num}. Abteilung")
            run.font.size = Pt(11)
            run.font.bold = True
            p_header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Linksbündig
            
            # Merge über alle Spalten
            cell_0 = header_cells[0]
            for i in range(1, 6):
                header_cells[i].paragraphs[0].clear()
                cell_0.merge(header_cells[i])
            
            # Grauer Hintergrund für Abteilungsheader
            tcPr = cell_0._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'E6E6E6')  # Hellgrau
            tcPr.append(shd)
        
        # Erste Startzeit pro Gruppe bestimmen
        group_start_time = None
        for starter in group_starters:
            start_time_candidate = starter.get("startTime")
            if start_time_candidate:
                group_start_time = start_time_candidate
                break
        # Starter der Gruppe verarbeiten
        for starter_idx, starter in enumerate(group_starters):
            row = table.add_row()
            row_cells = row.cells

            # Zebra: 0=weiß, 1=grau
            if row_counter % 2 == 1:
                set_light_gray_background(row)
            row_counter += 1
            
            # Prevent this row from breaking across pages and set compact height
            _set_row_keep_together(row)
            
            # Set row height to be more compact
            tr = row._tr
            trPr = tr.trPr
            if trPr is None:
                trPr = OxmlElement('w:trPr')
                tr.insert(0, trPr)
            
            # Set exact row height for compactness
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '280')  # Kompakte Zeilenhöhe in twips
            trHeight.set(qn('w:hRule'), 'exact')
            trPr.append(trHeight)
            
            # Column 0: Start number (NUR NUMMER)
            start_num = starter.get("startNumber", "")
            
            p0 = row_cells[0].paragraphs[0]
            p0.clear()
            p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            is_withdrawn = starter.get("withdrawn", False)
            is_hors_concours = starter.get("horsConcours", False)

            if start_num:
                run = p0.add_run(str(start_num))
                run.font.bold = True
                run.font.size = Pt(9)
                if is_withdrawn:
                    run.font.strike = True
            if is_hors_concours:
                p0.add_run("\n")
                ak_run = p0.add_run("AK")
                ak_run.font.size = Pt(7)
                ak_run.font.bold = False
            
            _set_compact_cell_format(p0)

            # Column 1: Zeit (immer anzeigen wenn keine Abteilungen, sonst nur erste Zeit pro Gruppe)
            p1 = row_cells[1].paragraphs[0]
            p1.clear()
            p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Zeit anzeigen: Entweder bei jedem Starter (wenn keine Abteilungen) oder nur beim ersten der Gruppe
            should_show_time = False
            if len(starters_by_group) == 1:  # Keine Abteilungen - zeige immer die Zeit
                individual_start_time = starter.get("startTime")
                if individual_start_time:
                    start_time_str = _format_time(individual_start_time)
                    should_show_time = True
            else:  # Abteilungen vorhanden - nur erste Zeit pro Gruppe
                if starter_idx == 0 and group_start_time:
                    start_time_str = _format_time(group_start_time)
                    should_show_time = True
            
            if should_show_time and start_time_str:
                run = p1.add_run(start_time_str)
                run.font.size = Pt(9)  # Zeit-Schriftgröße konsistent
            
            _set_compact_cell_format(p1)

            # Column 2: KoNr.
            horses = starter.get("horses", [])
            p2 = row_cells[2].paragraphs[0]
            p2.clear()
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            if horses:
                horse = horses[0]
                cno = _safe_get(horse, "cno", "")
                if cno:
                    run = p2.add_run(str(cno))
                    run.font.size = Pt(9)
            
            _set_compact_cell_format(p2)

            # Column 3: Reiter (NUR NAME - KOMPAKT, ABER BREITER)
            athlete = starter.get("athlete", {})
            rider_name = _safe_get(athlete, "name", "")

            club = athlete.get("club", "") if athlete else ""
            nationality = athlete.get("nation", "") if athlete else ""
            club_or_country = ""
            if nationality and nationality.upper() != "GER":
                if not club or club.strip() == "" or club.strip().upper() == "GASTLIZENZ GER":
                    club_or_country = get_country_name(nationality)
                else:
                    club_or_country = club
            elif club:
                club_or_country = club

            p3 = row_cells[3].paragraphs[0]
            p3.clear()
            p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if rider_name:
                run = p3.add_run(rider_name)
                run.font.bold = True
                run.font.size = Pt(9)
            if club_or_country:
                p3.add_run("\n")
                run_club = p3.add_run(club_or_country)
                run_club.font.size = Pt(7)
                run_club.font.bold = False
            
            _set_compact_cell_format(p3)

            # Column 4: Pferd (NUR NAME - KOMPAKT, ABER BREITER)
            p4 = row_cells[4].paragraphs[0]
            p4.clear()
            p4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            
            if horses:
                horse = horses[0]
                horse_name = _safe_get(horse, "name", "")
                
                # Nur Pferdename (KEINE Großschreibung)
                if horse_name:
                    run = p4.add_run(horse_name)
                    run.font.bold = True
                    run.font.size = Pt(9)
            
            _set_compact_cell_format(p4)

            # Column 5: Ergeb. (BREITER für Ergebnisse)
            p5 = row_cells[5].paragraphs[0]
            p5.clear()
            p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            _set_compact_cell_format(p5)

            # Handle withdrawn entries - nur durchgestrichen, Zebra-Farbe bleibt
            if is_withdrawn:
                for cell in row_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.strike = True

            # Check for breaks
            try:
                starter_num = int(start_num)
                if starter_num in breaks_map:
                    break_info = breaks_map[starter_num]
                    
                    break_row = table.add_row()
                    break_cells = break_row.cells
                    _set_row_keep_together(break_row)
                    
                    total_seconds = int(break_info.get("totalSeconds", 0) or 0)
                    info_text = break_info.get("informationText", "")
                    
                    if total_seconds > 0:
                        if total_seconds >= 3600:
                            hours = total_seconds // 3600
                            mins = (total_seconds % 3600) // 60
                            time_text = f"{hours} Std. {mins:02d} Min. Pause"
                        else:
                            mins = total_seconds // 60
                            time_text = f"{mins} Minuten Pause"
                        
                        break_text = f"{time_text} – {info_text}" if info_text else time_text
                    else:
                        break_text = info_text or "Pause"
                    
                    p_break = break_cells[0].paragraphs[0]
                    p_break.clear()
                    run = p_break.add_run(break_text)
                    run.font.size = Pt(10)
                    run.font.bold = True
                    p_break.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    p_break.paragraph_format.space_before = Pt(5)
                    p_break.paragraph_format.space_after  = Pt(5)
                    
                    # Kompakte Formatierung auch für Pausen-Text
                    _set_compact_cell_format(p_break)
                    
                    # Clear other cells and merge (6 columns for kompakt)
                    for i in range(1, 6):
                        break_cells[i].paragraphs[0].clear()
                    
                    cell_0 = break_cells[0]
                    for i in range(1, 6):
                        cell_0.merge(break_cells[i])
                    # Vertikale Zentrierung
                    tc_pr_b = cell_0._tc.get_or_add_tcPr()
                    v_align = OxmlElement('w:vAlign')
                    v_align.set(qn('w:val'), 'center')
                    tc_pr_b.append(v_align)
                    break_is_gray = (row_counter % 2 == 1)
                    row_counter += 1
                    if break_is_gray:
                        set_light_gray_background(break_row)
                    
            except (ValueError, TypeError):
                pass

    _set_table_borders(table)
    _add_sponsorship_footer(doc, show_sponsor_bar=show_sponsor_bar, show_banner=show_banner)

    # Add judges section - WICHTIG: Mit korrekter Reihenfolge E H C M B
    judges = starterlist.get("judges", [])
    if judges:
        doc.add_paragraph()

        # Aufgabe als Überschrift, nur wenn vorhanden
        dressage_tests = starterlist.get("dressageTests", [])
        test_name = dressage_tests[0].get("name", "") if dressage_tests else ""
        if test_name:
            p_aufgabe = doc.add_paragraph()
            p_aufgabe.add_run(f"Aufgabe: {test_name}").font.bold = True
            p_aufgabe.runs[0].font.size = Pt(11)
            p_aufgabe.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Get judge data for display - KORREKTE REIHENFOLGE E H C M B
        judge_display_data = _get_judge_data_for_display_kompakt(judges)
        
        if judge_display_data:
            judges_table = doc.add_table(rows=1 + len(judge_display_data), cols=2)
            
            # Try to set table style, but continue without if not available
            try:
                judges_table.style = 'Table Grid'
                print("WORD DEBUG: Applied Table Grid style to judges table")
            except Exception as e:
                print(f"WORD DEBUG: Table Grid style not available for judges table: {e}, continuing without style")
            
            judges_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Header row
            hdr_cells = judges_table.rows[0].cells
            headers = ["Pos", "Richter"]
            
            for i, header_text in enumerate(headers):
                p = hdr_cells[i].paragraphs[0]
                p.clear()
                run = p.add_run(header_text)
                run.font.bold = True
                run.font.size = Pt(9)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT if header_text == "Richter" else WD_PARAGRAPH_ALIGNMENT.CENTER
                _set_compact_cell_format(p)
            
            # Apply gray background to judges header row
            _set_header_gray_background(judges_table.rows[0])
            
            # Judge rows - KORREKTE REIHENFOLGE E H C M B
            for i, judge_data in enumerate(judge_display_data):
                row = judges_table.rows[i + 1]
                cells = row.cells
                _set_row_height_auto(row)
                
                # Position
                p = cells[0].paragraphs[0]
                p.clear()
                run = p.add_run(judge_data.get("pos_label", ""))
                run.font.size = Pt(9)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                _set_compact_cell_format(p)
                
                # Name
                p = cells[1].paragraphs[0]
                p.clear()
                run = p.add_run(judge_data.get("name", ""))
                run.font.size = Pt(9)
                _set_compact_cell_format(p)

            _set_table_borders(judges_table)
            _set_judges_table_widths(judges_table)

    # Save document
    print("WORD DEBUG: Starting document save")
    try:
        doc.save(output_path)
        print(f"WORD STANDARD LOGO DEBUG: Document saved to {output_path}")
        return output_path
    except Exception as e:
        print(f"WORD DEBUG: Error saving document: {e}")
        alt_output_path = filename
        try:
            doc.save(alt_output_path)
            print(f"WORD STANDARD LOGO DEBUG: Document saved to alternative location: {alt_output_path}")
            return alt_output_path
        except Exception as e2:
            print(f"WORD DEBUG: Failed to save to alternative location: {e2}")
            raise