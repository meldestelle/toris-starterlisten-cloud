# -*- coding: utf-8 -*-
# templates/stream/stream_int.py - International Design - EXAKT wie PDF
import os
import json
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT

def _set_odd_even_headers(doc):
    """Setzt 'Gerade & ungerade Seiten unterschiedlich' in Word-Einstellungen"""
    settings = doc.settings.element
    # Prüfen ob bereits vorhanden
    if settings.find(qn('w:evenAndOddHeaders')) is None:
        even_odd = OxmlElement('w:evenAndOddHeaders')
        settings.insert(0, even_odd)


OUTPUT_DIR = "Ausgabe"

def _ensure_output_dir():
    """Stellt sicher, dass das Ausgabe-Verzeichnis existiert"""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

def safe_get(dictionary, key, default=""):
    """Sicher Dictionary-Werte abrufen"""
    if not dictionary:
        return default
    if isinstance(dictionary, dict):
        return dictionary.get(key, default)
    return default

def format_header_datetime(time_iso):
    """Zeit formatieren für Header"""
    if not time_iso:
        return ""
    try:
        dt = datetime.fromisoformat(str(time_iso).replace("Z", ""))
        weekday_map = {
            "Monday": "Montag", "Tuesday": "Dienstag", "Wednesday": "Mittwoch",
            "Thursday": "Donnerstag", "Friday": "Freitag", "Saturday": "Samstag",
            "Sunday": "Sonntag"
        }
        month_map = {
            "January": "Januar", "February": "Februar", "March": "März", "April": "April",
            "May": "Mai", "June": "Juni", "July": "Juli", "August": "August",
            "September": "September", "October": "Oktober", "November": "November", "December": "Dezember"
        }
        weekday_en = dt.strftime("%A")
        month_en = dt.strftime("%B")
        weekday_de = weekday_map.get(weekday_en, weekday_en)
        month_de = month_map.get(month_en, month_en)
        return f"{weekday_de}, {dt.day}. {month_de} {dt.year} {dt.strftime('%H:%M')} Uhr"
    except:
        return str(time_iso)

def format_time(time_iso):
    """Zeit formatieren MIT SEKUNDEN (HH:MM:SS)"""
    if not time_iso:
        return ""
    try:
        dt = datetime.fromisoformat(str(time_iso).replace("Z", ""))
        return dt.strftime("%H:%M:%S")  # MIT SEKUNDEN!
    except:
        time_str = str(time_iso)
        return time_str[:8] if len(time_str) >= 8 else time_str

def get_flag_symbol(nationality_str):
    """Gibt Flaggen-Symbol als Text zurück (wie im PDF)"""
    if not nationality_str:
        return ""
    
    code = get_nationality_code(nationality_str)
    
    # Mapping zu Flaggen-Boxen (wie im PDF: kleine Kästchen mit Farbcode)
    # Wir verwenden Emoji-Flags die in Word besser funktionieren
    flag_map = {
        "GER": "🇩🇪", "NED": "🇳🇱", "BEL": "🇧🇪", "FRA": "🇫🇷",
        "AUT": "🇦🇹", "SUI": "🇨🇭", "POL": "🇵🇱", "LUX": "🇱🇺",
        "ESP": "🇪🇸", "ITA": "🇮🇹", "GBR": "🇬🇧", "USA": "🇺🇸",
        "CAN": "🇨🇦", "AUS": "🇦🇺", "SWE": "🇸🇪", "DEN": "🇩🇰",
        "NOR": "🇳🇴", "FIN": "🇫🇮", "RUS": "🇷🇺", "UKR": "🇺🇦",
        "HUN": "🇭🇺", "CZE": "🇨🇿", "IRL": "🇮🇪", "POR": "🇵🇹"
    }
    
    return flag_map.get(code, "")

def get_nationality_code(nationality_str):
    """
    Konvertiert IOC-Codes zu ISO-Codes für Flaggen-Dateien.
    WICHTIG: GER bleibt GER (nicht DEU), weil Flaggen als GER.png benannt sind!
    """
    if not nationality_str:
        return ""
    
    code = str(nationality_str).strip().upper()
    
    # Mapping IOC → Flaggen-Code (meist IOC = Flaggenname)
    ioc_to_flag = {
        # A
        "AFG": "AFG",  # Afghanistan
        "ALB": "ALB",  # Albania
        "ALG": "ALG",  # Algeria
        "AND": "AND",  # Andorra
        "ANG": "ANG",  # Angola
        "ARG": "ARG",  # Argentina
        "ARM": "ARM",  # Armenia
        "ARU": "ARU",  # Aruba
        "ASA": "ASA",  # American Samoa
        "AUS": "AUS",  # Australia
        "AUT": "AUT",  # Austria
        "AZE": "AZE",  # Azerbaijan
        
        # B
        "BAH": "BAH",  # Bahamas
        "BAN": "BAN",  # Bangladesh
        "BAR": "BAR",  # Barbados
        "BDI": "BDI",  # Burundi
        "BEL": "BEL",  # Belgium
        "BEN": "BEN",  # Benin
        "BER": "BER",  # Bermuda
        "BHU": "BHU",  # Bhutan
        "BIH": "BIH",  # Bosnia and Herzegovina
        "BIZ": "BIZ",  # Belize
        "BLR": "BLR",  # Belarus
        "BOL": "BOL",  # Bolivia
        "BOT": "BOT",  # Botswana
        "BRA": "BRA",  # Brazil
        "BRN": "BRN",  # Bahrain
        "BRU": "BRU",  # Brunei
        "BUL": "BUL",  # Bulgaria
        "BUR": "BUR",  # Burkina Faso
        
        # C
        "CAF": "CAF",  # Central African Republic
        "CAM": "CAM",  # Cambodia
        "CAN": "CAN",  # Canada
        "CAY": "CAY",  # Cayman Islands
        "CGO": "CGO",  # Republic of the Congo
        "CHA": "CHA",  # Chad
        "CHI": "CHI",  # Chile
        "CHN": "CHN",  # China
        "CIV": "CIV",  # Côte d'Ivoire
        "CMR": "CMR",  # Cameroon
        "COD": "COD",  # Democratic Republic of the Congo
        "COK": "COK",  # Cook Islands
        "COL": "COL",  # Colombia
        "COM": "COM",  # Comoros
        "CPV": "CPV",  # Cape Verde
        "CRC": "CRC",  # Costa Rica
        "CRO": "CRO",  # Croatia
        "CUB": "CUB",  # Cuba
        "CYP": "CYP",  # Cyprus
        "CZE": "CZE",  # Czech Republic
        
        # D
        "DEN": "DEN",  # Denmark
        "DJI": "DJI",  # Djibouti
        "DMA": "DMA",  # Dominica
        "DOM": "DOM",  # Dominican Republic
        
        # E
        "ECU": "ECU",  # Ecuador
        "EGY": "EGY",  # Egypt
        "ERI": "ERI",  # Eritrea
        "ESA": "ESA",  # El Salvador
        "ESP": "ESP",  # Spain
        "EST": "EST",  # Estonia
        "ETH": "ETH",  # Ethiopia
        
        # F
        "FIJ": "FIJ",  # Fiji
        "FIN": "FIN",  # Finland
        "FRA": "FRA",  # France
        
        # G
        "GAB": "GAB",  # Gabon
        "GAM": "GAM",  # Gambia
        "GBR": "GBR",  # Great Britain
        "GBS": "GBS",  # Guinea-Bissau
        "GEO": "GEO",  # Georgia
        "GEQ": "GEQ",  # Equatorial Guinea
        "GER": "GER",  # Germany - BLEIBT GER!
        "GHA": "GHA",  # Ghana
        "GRE": "GRE",  # Greece
        "GRN": "GRN",  # Grenada
        "GUA": "GUA",  # Guatemala
        "GUI": "GUI",  # Guinea
        "GUM": "GUM",  # Guam
        "GUY": "GUY",  # Guyana
        
        # H
        "HAI": "HAI",  # Haiti
        "HKG": "HKG",  # Hong Kong
        "HON": "HON",  # Honduras
        "HUN": "HUN",  # Hungary
        
        # I
        "INA": "INA",  # Indonesia
        "IND": "IND",  # India
        "IRI": "IRI",  # Iran
        "IRL": "IRL",  # Ireland
        "IRQ": "IRQ",  # Iraq
        "ISL": "ISL",  # Iceland
        "ISR": "ISR",  # Israel
        "ISV": "ISV",  # Virgin Islands
        "ITA": "ITA",  # Italy
        
        # J
        "JAM": "JAM",  # Jamaica
        "JOR": "JOR",  # Jordan
        "JPN": "JPN",  # Japan
        
        # K
        "KAZ": "KAZ",  # Kazakhstan
        "KEN": "KEN",  # Kenya
        "KGZ": "KGZ",  # Kyrgyzstan
        "KIR": "KIR",  # Kiribati
        "KOR": "KOR",  # South Korea
        "KOS": "KOS",  # Kosovo
        "KSA": "KSA",  # Saudi Arabia
        "KUW": "KUW",  # Kuwait
        
        # L
        "LAO": "LAO",  # Laos
        "LAT": "LAT",  # Latvia
        "LBA": "LBA",  # Libya
        "LBR": "LBR",  # Liberia
        "LCA": "LCA",  # Saint Lucia
        "LES": "LES",  # Lesotho
        "LIB": "LIB",  # Lebanon
        "LIE": "LIE",  # Liechtenstein
        "LTU": "LTU",  # Lithuania
        "LUX": "LUX",  # Luxembourg
        
        # M
        "MAD": "MAD",  # Madagascar
        "MAR": "MAR",  # Morocco
        "MAS": "MAS",  # Malaysia
        "MAW": "MAW",  # Malawi
        "MDA": "MDA",  # Moldova
        "MDV": "MDV",  # Maldives
        "MEX": "MEX",  # Mexico
        "MGL": "MGL",  # Mongolia
        "MHL": "MHL",  # Marshall Islands
        "MKD": "MKD",  # North Macedonia
        "MLI": "MLI",  # Mali
        "MLT": "MLT",  # Malta
        "MNE": "MNE",  # Montenegro
        "MON": "MON",  # Monaco
        "MOZ": "MOZ",  # Mozambique
        "MRI": "MRI",  # Mauritius
        "MTN": "MTN",  # Mauritania
        "MYA": "MYA",  # Myanmar
        
        # N
        "NAM": "NAM",  # Namibia
        "NCA": "NCA",  # Nicaragua
        "NED": "NED",  # Netherlands
        "NEP": "NEP",  # Nepal
        "NGR": "NGR",  # Nigeria
        "NIG": "NIG",  # Niger
        "NOR": "NOR",  # Norway
        "NRU": "NRU",  # Nauru
        "NZL": "NZL",  # New Zealand
        
        # O
        "OMA": "OMA",  # Oman
        
        # P
        "PAK": "PAK",  # Pakistan
        "PAN": "PAN",  # Panama
        "PAR": "PAR",  # Paraguay
        "PER": "PER",  # Peru
        "PHI": "PHI",  # Philippines
        "PLE": "PLE",  # Palestine
        "PLW": "PLW",  # Palau
        "PNG": "PNG",  # Papua New Guinea
        "POL": "POL",  # Poland
        "POR": "POR",  # Portugal
        "PRK": "PRK",  # North Korea
        "PUR": "PUR",  # Puerto Rico
        
        # Q
        "QAT": "QAT",  # Qatar
        
        # R
        "ROU": "ROU",  # Romania
        "RSA": "RSA",  # South Africa
        "RUS": "RUS",  # Russia
        "RWA": "RWA",  # Rwanda
        
        # S
        "SAM": "SAM",  # Samoa
        "SEN": "SEN",  # Senegal
        "SEY": "SEY",  # Seychelles
        "SGP": "SGP",  # Singapore
        "SKN": "SKN",  # Saint Kitts and Nevis
        "SLE": "SLE",  # Sierra Leone
        "SLO": "SLO",  # Slovenia
        "SMR": "SMR",  # San Marino
        "SOL": "SOL",  # Solomon Islands
        "SOM": "SOM",  # Somalia
        "SRB": "SRB",  # Serbia
        "SRI": "SRI",  # Sri Lanka
        "STP": "STP",  # São Tomé and Príncipe
        "SUD": "SUD",  # Sudan
        "SUI": "SUI",  # Switzerland
        "SUR": "SUR",  # Suriname
        "SVK": "SVK",  # Slovakia
        "SWE": "SWE",  # Sweden
        "SWZ": "SWZ",  # Eswatini
        "SYR": "SYR",  # Syria
        
        # T
        "TAN": "TAN",  # Tanzania
        "TGA": "TGA",  # Tonga
        "THA": "THA",  # Thailand
        "TJK": "TJK",  # Tajikistan
        "TKM": "TKM",  # Turkmenistan
        "TLS": "TLS",  # Timor-Leste
        "TOG": "TOG",  # Togo
        "TPE": "TPE",  # Chinese Taipei
        "TTO": "TTO",  # Trinidad and Tobago
        "TUN": "TUN",  # Tunisia
        "TUR": "TUR",  # Turkey
        "TUV": "TUV",  # Tuvalu
        
        # U
        "UAE": "UAE",  # United Arab Emirates
        "UGA": "UGA",  # Uganda
        "UKR": "UKR",  # Ukraine
        "URU": "URU",  # Uruguay
        "USA": "USA",  # United States
        "UZB": "UZB",  # Uzbekistan
        
        # V
        "VAN": "VAN",  # Vanuatu
        "VEN": "VEN",  # Venezuela
        "VIE": "VIE",  # Vietnam
        "VIN": "VIN",  # Saint Vincent and the Grenadines
        
        # Y
        "YEM": "YEM",  # Yemen
        
        # Z
        "ZAM": "ZAM",  # Zambia
        "ZIM": "ZIM",  # Zimbabwe
        
        # Vollständige Ländernamen (fallback)
        "GERMANY": "GER",
        "DEUTSCHLAND": "GER",
        "NETHERLANDS": "NED",
        "NIEDERLANDE": "NED",
        "BELGIUM": "BEL",
        "BELGIEN": "BEL",
        "FRANCE": "FRA",
        "FRANKREICH": "FRA",
        "AUSTRIA": "AUT",
        "ÖSTERREICH": "AUT",
        "OESTERREICH": "AUT",
        "SWITZERLAND": "SUI",
        "SCHWEIZ": "SUI",
        "POLAND": "POL",
        "POLEN": "POL",
        "LUXEMBOURG": "LUX",
        "LUXEMBURG": "LUX",
        "SPAIN": "ESP",
        "SPANIEN": "ESP",
        "ITALY": "ITA",
        "ITALIEN": "ITA",
        "GREAT BRITAIN": "GBR",
        "GROSSBRITANNIEN": "GBR",
        "UNITED KINGDOM": "GBR",
    }
    
    # Direkte Rückgabe wenn Code im Mapping
    if code in ioc_to_flag:
        return ioc_to_flag[code]
    
    # Fallback: Original zurückgeben
    return code


def get_country_name_german(code):
    """Gibt den deutschen Ländernamen zurück"""
    if not code:
        return ""
    
    code = str(code).strip().upper()
    
    names = {
        # A
        "AFG": "Afghanistan",
        "ALB": "Albanien",
        "ALG": "Algerien",
        "AND": "Andorra",
        "ANG": "Angola",
        "ARG": "Argentinien",
        "ARM": "Armenien",
        "ARU": "Aruba",
        "ASA": "Amerikanisch-Samoa",
        "AUS": "Australien",
        "AUT": "Österreich",
        "AZE": "Aserbaidschan",
        
        # B
        "BAH": "Bahamas",
        "BAN": "Bangladesch",
        "BAR": "Barbados",
        "BDI": "Burundi",
        "BEL": "Belgien",
        "BEN": "Benin",
        "BER": "Bermuda",
        "BHU": "Bhutan",
        "BIH": "Bosnien und Herzegowina",
        "BIZ": "Belize",
        "BLR": "Belarus",
        "BOL": "Bolivien",
        "BOT": "Botswana",
        "BRA": "Brasilien",
        "BRN": "Bahrain",
        "BRU": "Brunei",
        "BUL": "Bulgarien",
        "BUR": "Burkina Faso",
        
        # C
        "CAF": "Zentralafrikanische Republik",
        "CAM": "Kambodscha",
        "CAN": "Kanada",
        "CAY": "Kaimaninseln",
        "CGO": "Republik Kongo",
        "CHA": "Tschad",
        "CHI": "Chile",
        "CHN": "China",
        "CIV": "Elfenbeinküste",
        "CMR": "Kamerun",
        "COD": "Demokratische Republik Kongo",
        "COK": "Cookinseln",
        "COL": "Kolumbien",
        "COM": "Komoren",
        "CPV": "Kap Verde",
        "CRC": "Costa Rica",
        "CRO": "Kroatien",
        "CUB": "Kuba",
        "CYP": "Zypern",
        "CZE": "Tschechien",
        
        # D
        "DEN": "Dänemark",
        "DJI": "Dschibuti",
        "DMA": "Dominica",
        "DOM": "Dominikanische Republik",
        
        # E
        "ECU": "Ecuador",
        "EGY": "Ägypten",
        "ERI": "Eritrea",
        "ESA": "El Salvador",
        "ESP": "Spanien",
        "EST": "Estland",
        "ETH": "Äthiopien",
        
        # F
        "FIJ": "Fidschi",
        "FIN": "Finnland",
        "FRA": "Frankreich",
        
        # G
        "GAB": "Gabun",
        "GAM": "Gambia",
        "GBR": "Großbritannien",
        "GBS": "Guinea-Bissau",
        "GEO": "Georgien",
        "GEQ": "Äquatorialguinea",
        "GER": "Deutschland",
        "GHA": "Ghana",
        "GRE": "Griechenland",
        "GRN": "Grenada",
        "GUA": "Guatemala",
        "GUI": "Guinea",
        "GUM": "Guam",
        "GUY": "Guyana",
        
        # H
        "HAI": "Haiti",
        "HKG": "Hongkong",
        "HON": "Honduras",
        "HUN": "Ungarn",
        
        # I
        "INA": "Indonesien",
        "IND": "Indien",
        "IRI": "Iran",
        "IRL": "Irland",
        "IRQ": "Irak",
        "ISL": "Island",
        "ISR": "Israel",
        "ISV": "Amerikanische Jungferninseln",
        "ITA": "Italien",
        
        # J
        "JAM": "Jamaika",
        "JOR": "Jordanien",
        "JPN": "Japan",
        
        # K
        "KAZ": "Kasachstan",
        "KEN": "Kenia",
        "KGZ": "Kirgisistan",
        "KIR": "Kiribati",
        "KOR": "Südkorea",
        "KOS": "Kosovo",
        "KSA": "Saudi-Arabien",
        "KUW": "Kuwait",
        
        # L
        "LAO": "Laos",
        "LAT": "Lettland",
        "LBA": "Libyen",
        "LBR": "Liberia",
        "LCA": "St. Lucia",
        "LES": "Lesotho",
        "LIB": "Libanon",
        "LIE": "Liechtenstein",
        "LTU": "Litauen",
        "LUX": "Luxemburg",
        
        # M
        "MAD": "Madagaskar",
        "MAR": "Marokko",
        "MAS": "Malaysia",
        "MAW": "Malawi",
        "MDA": "Moldau",
        "MDV": "Malediven",
        "MEX": "Mexiko",
        "MGL": "Mongolei",
        "MHL": "Marshallinseln",
        "MKD": "Nordmazedonien",
        "MLI": "Mali",
        "MLT": "Malta",
        "MNE": "Montenegro",
        "MON": "Monaco",
        "MOZ": "Mosambik",
        "MRI": "Mauritius",
        "MTN": "Mauretanien",
        "MYA": "Myanmar",
        
        # N
        "NAM": "Namibia",
        "NCA": "Nicaragua",
        "NED": "Niederlande",
        "NEP": "Nepal",
        "NGR": "Nigeria",
        "NIG": "Niger",
        "NOR": "Norwegen",
        "NRU": "Nauru",
        "NZL": "Neuseeland",
        
        # O
        "OMA": "Oman",
        
        # P
        "PAK": "Pakistan",
        "PAN": "Panama",
        "PAR": "Paraguay",
        "PER": "Peru",
        "PHI": "Philippinen",
        "PLE": "Palästina",
        "PLW": "Palau",
        "PNG": "Papua-Neuguinea",
        "POL": "Polen",
        "POR": "Portugal",
        "PRK": "Nordkorea",
        "PUR": "Puerto Rico",
        
        # Q
        "QAT": "Katar",
        
        # R
        "ROU": "Rumänien",
        "RSA": "Südafrika",
        "RUS": "Russland",
        "RWA": "Ruanda",
        
        # S
        "SAM": "Samoa",
        "SEN": "Senegal",
        "SEY": "Seychellen",
        "SGP": "Singapur",
        "SKN": "St. Kitts und Nevis",
        "SLE": "Sierra Leone",
        "SLO": "Slowenien",
        "SMR": "San Marino",
        "SOL": "Salomonen",
        "SOM": "Somalia",
        "SRB": "Serbien",
        "SRI": "Sri Lanka",
        "STP": "São Tomé und Príncipe",
        "SUD": "Sudan",
        "SUI": "Schweiz",
        "SUR": "Suriname",
        "SVK": "Slowakei",
        "SWE": "Schweden",
        "SWZ": "Eswatini",
        "SYR": "Syrien",
        
        # T
        "TAN": "Tansania",
        "TGA": "Tonga",
        "THA": "Thailand",
        "TJK": "Tadschikistan",
        "TKM": "Turkmenistan",
        "TLS": "Timor-Leste",
        "TOG": "Togo",
        "TPE": "Chinesisch Taipeh",
        "TTO": "Trinidad und Tobago",
        "TUN": "Tunesien",
        "TUR": "Türkei",
        "TUV": "Tuvalu",
        
        # U
        "UAE": "Vereinigte Arabische Emirate",
        "UGA": "Uganda",
        "UKR": "Ukraine",
        "URU": "Uruguay",
        "USA": "Vereinigte Staaten",
        "UZB": "Usbekistan",
        
        # V
        "VAN": "Vanuatu",
        "VEN": "Venezuela",
        "VIE": "Vietnam",
        "VIN": "St. Vincent und die Grenadinen",
        
        # Y
        "YEM": "Jemen",
        
        # Z
        "ZAM": "Sambia",
        "ZIM": "Simbabwe",
    }
    
    return names.get(code, code)


def get_country_name_english(code):
    """Gibt den englischen Ländernamen zurück"""
    if not code:
        return ""
    
    code = str(code).strip().upper()
    
    names = {
        # A
        "AFG": "Afghanistan",
        "ALB": "Albania",
        "ALG": "Algeria",
        "AND": "Andorra",
        "ANG": "Angola",
        "ARG": "Argentina",
        "ARM": "Armenia",
        "ARU": "Aruba",
        "ASA": "American Samoa",
        "AUS": "Australia",
        "AUT": "Austria",
        "AZE": "Azerbaijan",
        
        # B
        "BAH": "Bahamas",
        "BAN": "Bangladesh",
        "BAR": "Barbados",
        "BDI": "Burundi",
        "BEL": "Belgium",
        "BEN": "Benin",
        "BER": "Bermuda",
        "BHU": "Bhutan",
        "BIH": "Bosnia and Herzegovina",
        "BIZ": "Belize",
        "BLR": "Belarus",
        "BOL": "Bolivia",
        "BOT": "Botswana",
        "BRA": "Brazil",
        "BRN": "Bahrain",
        "BRU": "Brunei",
        "BUL": "Bulgaria",
        "BUR": "Burkina Faso",
        
        # C
        "CAF": "Central African Republic",
        "CAM": "Cambodia",
        "CAN": "Canada",
        "CAY": "Cayman Islands",
        "CGO": "Republic of the Congo",
        "CHA": "Chad",
        "CHI": "Chile",
        "CHN": "China",
        "CIV": "Ivory Coast",
        "CMR": "Cameroon",
        "COD": "Democratic Republic of the Congo",
        "COK": "Cook Islands",
        "COL": "Colombia",
        "COM": "Comoros",
        "CPV": "Cape Verde",
        "CRC": "Costa Rica",
        "CRO": "Croatia",
        "CUB": "Cuba",
        "CYP": "Cyprus",
        "CZE": "Czech Republic",
        
        # D
        "DEN": "Denmark",
        "DJI": "Djibouti",
        "DMA": "Dominica",
        "DOM": "Dominican Republic",
        
        # E
        "ECU": "Ecuador",
        "EGY": "Egypt",
        "ERI": "Eritrea",
        "ESA": "El Salvador",
        "ESP": "Spain",
        "EST": "Estonia",
        "ETH": "Ethiopia",
        
        # F
        "FIJ": "Fiji",
        "FIN": "Finland",
        "FRA": "France",
        
        # G
        "GAB": "Gabon",
        "GAM": "Gambia",
        "GBR": "Great Britain",
        "GBS": "Guinea-Bissau",
        "GEO": "Georgia",
        "GEQ": "Equatorial Guinea",
        "GER": "Germany",
        "GHA": "Ghana",
        "GRE": "Greece",
        "GRN": "Grenada",
        "GUA": "Guatemala",
        "GUI": "Guinea",
        "GUM": "Guam",
        "GUY": "Guyana",
        
        # H
        "HAI": "Haiti",
        "HKG": "Hong Kong",
        "HON": "Honduras",
        "HUN": "Hungary",
        
        # I
        "INA": "Indonesia",
        "IND": "India",
        "IRI": "Iran",
        "IRL": "Ireland",
        "IRQ": "Iraq",
        "ISL": "Iceland",
        "ISR": "Israel",
        "ISV": "Virgin Islands",
        "ITA": "Italy",
        
        # J
        "JAM": "Jamaica",
        "JOR": "Jordan",
        "JPN": "Japan",
        
        # K
        "KAZ": "Kazakhstan",
        "KEN": "Kenya",
        "KGZ": "Kyrgyzstan",
        "KIR": "Kiribati",
        "KOR": "South Korea",
        "KOS": "Kosovo",
        "KSA": "Saudi Arabia",
        "KUW": "Kuwait",
        
        # L
        "LAO": "Laos",
        "LAT": "Latvia",
        "LBA": "Libya",
        "LBR": "Liberia",
        "LCA": "Saint Lucia",
        "LES": "Lesotho",
        "LIB": "Lebanon",
        "LIE": "Liechtenstein",
        "LTU": "Lithuania",
        "LUX": "Luxembourg",
        
        # M
        "MAD": "Madagascar",
        "MAR": "Morocco",
        "MAS": "Malaysia",
        "MAW": "Malawi",
        "MDA": "Moldova",
        "MDV": "Maldives",
        "MEX": "Mexico",
        "MGL": "Mongolia",
        "MHL": "Marshall Islands",
        "MKD": "North Macedonia",
        "MLI": "Mali",
        "MLT": "Malta",
        "MNE": "Montenegro",
        "MON": "Monaco",
        "MOZ": "Mozambique",
        "MRI": "Mauritius",
        "MTN": "Mauritania",
        "MYA": "Myanmar",
        
        # N
        "NAM": "Namibia",
        "NCA": "Nicaragua",
        "NED": "Netherlands",
        "NEP": "Nepal",
        "NGR": "Nigeria",
        "NIG": "Niger",
        "NOR": "Norway",
        "NRU": "Nauru",
        "NZL": "New Zealand",
        
        # O
        "OMA": "Oman",
        
        # P
        "PAK": "Pakistan",
        "PAN": "Panama",
        "PAR": "Paraguay",
        "PER": "Peru",
        "PHI": "Philippines",
        "PLE": "Palestine",
        "PLW": "Palau",
        "PNG": "Papua New Guinea",
        "POL": "Poland",
        "POR": "Portugal",
        "PRK": "North Korea",
        "PUR": "Puerto Rico",
        
        # Q
        "QAT": "Qatar",
        
        # R
        "ROU": "Romania",
        "RSA": "South Africa",
        "RUS": "Russia",
        "RWA": "Rwanda",
        
        # S
        "SAM": "Samoa",
        "SEN": "Senegal",
        "SEY": "Seychelles",
        "SGP": "Singapore",
        "SKN": "Saint Kitts and Nevis",
        "SLE": "Sierra Leone",
        "SLO": "Slovenia",
        "SMR": "San Marino",
        "SOL": "Solomon Islands",
        "SOM": "Somalia",
        "SRB": "Serbia",
        "SRI": "Sri Lanka",
        "STP": "São Tomé and Príncipe",
        "SUD": "Sudan",
        "SUI": "Switzerland",
        "SUR": "Suriname",
        "SVK": "Slovakia",
        "SWE": "Sweden",
        "SWZ": "Eswatini",
        "SYR": "Syria",
        
        # T
        "TAN": "Tanzania",
        "TGA": "Tonga",
        "THA": "Thailand",
        "TJK": "Tajikistan",
        "TKM": "Turkmenistan",
        "TLS": "Timor-Leste",
        "TOG": "Togo",
        "TPE": "Chinese Taipei",
        "TTO": "Trinidad and Tobago",
        "TUN": "Tunisia",
        "TUR": "Turkey",
        "TUV": "Tuvalu",
        
        # U
        "UAE": "United Arab Emirates",
        "UGA": "Uganda",
        "UKR": "Ukraine",
        "URU": "Uruguay",
        "USA": "United States",
        "UZB": "Uzbekistan",
        
        # V
        "VAN": "Vanuatu",
        "VEN": "Venezuela",
        "VIE": "Vietnam",
        "VIN": "Saint Vincent and the Grenadines",
        
        # Y
        "YEM": "Yemen",
        
        # Z
        "ZAM": "Zambia",
        "ZIM": "Zimbabwe",
    }
    
    return names.get(code, code)



def set_horizontal_borders_only(table):
    """NUR horizontale Borders - KEINE vertikalen!
    Entfernt AGGRESSIV alle Zell-Borders"""
    
    # 1. Table-Level: Nur horizontale Borders
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Entferne alle bestehenden Table-Borders
    tblBorders = tblPr.find(qn('w:tblBorders'))
    if tblBorders is not None:
        tblPr.remove(tblBorders)
    
    # Setze NUR horizontale Table-Borders
    tblBorders = OxmlElement('w:tblBorders')
    
    # Top Border
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:color'), '000000')
    tblBorders.append(top)
    
    # Bottom Border
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), '000000')
    tblBorders.append(bottom)
    
    # Inside Horizontal
    insideH = OxmlElement('w:insideH')
    insideH.set(qn('w:val'), 'single')
    insideH.set(qn('w:sz'), '4')
    insideH.set(qn('w:color'), '000000')
    tblBorders.append(insideH)
    
    # EXPLIZIT: Keine vertikalen Borders!
    for side in ['left', 'right', 'insideV']:
        elem = OxmlElement(f'w:{side}')
        elem.set(qn('w:val'), 'none')
        elem.set(qn('w:sz'), '0')
        tblBorders.append(elem)
    
    tblPr.append(tblBorders)
    
    # 2. Cell-Level: ENTFERNE ALLE CELL-BORDERS
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Entferne bestehende Cell-Borders
            existing_borders = tcPr.xpath('.//w:tcBorders')
            for border in existing_borders:
                tcPr.remove(border)
            
            # Setze explizit KEINE Borders auf Cell-Level
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                tcBorders.append(border)
            tcPr.append(tcBorders)

def set_column_widths(table):
    """Spaltenbreiten für PDF-Layout"""
    if not table.rows:
        return
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)
    
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    tblGrid = OxmlElement('w:tblGrid')
    # # | Time | CNO (20% schmaler) | Horse (10% schmaler) | Athlete (breiter) | Nat. (SCHMAL!)
    # Horse: 6500 - 10% = 5850
    # Athlete: 3680 + 650 = 4330
    widths = [600, 1100, 720, 5850, 4330, 500]
    
    for width in widths:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)
    
    if tblPr.getparent() is not None:
        tblPr.getparent().insert(tblPr.getparent().index(tblPr) + 1, tblGrid)
    else:
        tbl.insert(0, tblGrid)
    
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                tcPr = cell._tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(widths[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def set_row_height_auto(row, cant_split=False):
    """Set row height to auto, optionally prevent page breaks within row"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is not None:
        trPr.remove(trHeight)
    
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '0')
    trHeight.set(qn('w:hRule'), 'auto')
    trPr.append(trHeight)

    if cant_split:
        cantSplit = trPr.find(qn('w:cantSplit'))
        if cantSplit is None:
            cantSplit = OxmlElement('w:cantSplit')
            trPr.append(cantSplit)
        cantSplit.set(qn('w:val'), '1')

def set_cell_margins_tight(table):
    """Set tight cell margins"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
    
    tblCellMar = OxmlElement('w:tblCellMar')
    
    top = OxmlElement('w:top')
    top.set(qn('w:w'), '72')
    top.set(qn('w:type'), 'dxa')
    tblCellMar.append(top)
    
    left = OxmlElement('w:left')
    left.set(qn('w:w'), '108')
    left.set(qn('w:type'), 'dxa')
    tblCellMar.append(left)
    
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:w'), '72')
    bottom.set(qn('w:type'), 'dxa')
    tblCellMar.append(bottom)
    
    right = OxmlElement('w:right')
    right.set(qn('w:w'), '108')
    right.set(qn('w:type'), 'dxa')
    tblCellMar.append(right)
    
    tblPr.append(tblCellMar)

def optimize_paragraph_spacing(paragraph):
    """Optimize paragraph spacing"""
    pPr = paragraph._element.get_or_add_pPr()
    
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)
    
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    spacing.set(qn('w:line'), '240')
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def set_dark_header_background(row):
    """Dunkelgrauer/schwarzer Hintergrund für Header (wie im PDF)"""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '404040')  # Dunkelgrau
        tcPr.append(shd)

def set_light_gray_background(row):
    """Hellgraue Hintergrundfarbe (wie im PDF)"""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'E8E8E8')  # Hellgrau
        tcPr.append(shd)

def render(starterlist, filename):
    """Hauptfunktion - International Style MIT Logo und Header"""
    print("WORD_INT DEBUG: Starting render function")
    
    if starterlist is None:
        raise ValueError("starterlist is None")
    
    if isinstance(starterlist, str):
        try:
            starterlist = json.loads(starterlist)
        except Exception as e:
            raise ValueError(f"starterlist is a string but not valid JSON: {e}")
    
    _ensure_output_dir()
    # Druckoptionen auslesen
    print_options     = starterlist.get("printOptions", {})
    show_header       = print_options.get("show_header",      True)
    show_banner       = print_options.get("show_banner",      True)
    show_sponsor_bar  = print_options.get("show_sponsor_bar", True)
    show_title_opt    = print_options.get("show_title",       True)
    sponsor_top       = print_options.get("sponsor_top",      False)
    sponsor_bottom    = print_options.get("sponsor_bottom",   False)
    spacing_top_cm    = starterlist.get("spacingTopCm",   3.0)
    spacing_bottom_cm = starterlist.get("spacingBottomCm",2.0)
    
    # Output path - IMMER ins Ausgabe-Verzeichnis!
    output_path = os.path.join(OUTPUT_DIR, filename)
    print(f"WORD_INT DEBUG: Output path: {output_path}")
    
    # Neues Dokument erstellen
    doc = Document()

    # Seitenformat + Banner
    _section = doc.sections[0]
    _section.page_height = Inches(11.69)
    _section.page_width  = Inches(8.27)
    _section.left_margin  = Inches(0.5)
    _section.right_margin = Inches(0.5)
    _MIN_TOP    = Inches(10 / 25.4)
    _MIN_BOTTOM = Inches(0.26)
    _SP_TOP     = Inches(spacing_top_cm / 2.54)
    _SP_BOTTOM  = Inches(spacing_bottom_cm / 2.54)

    if sponsor_top or sponsor_bottom:
        _section.different_odd_and_even_pages_header_footer = True
        _set_odd_even_headers(doc)
    if sponsor_top:
        _section.header_distance = Inches(0.2)
        _section.top_margin      = _SP_TOP
    else:
        _section.header_distance = Inches(0.2)
        _section.top_margin      = _MIN_TOP

    if sponsor_bottom:
        _section.footer_distance = _SP_BOTTOM
        _section.bottom_margin   = _SP_BOTTOM
    else:
        _section.footer_distance = Inches(0.24)
        _section.bottom_margin   = _MIN_BOTTOM
    if show_banner:
        _bp = "logos/banner.png"
        if os.path.exists(_bp):
            try:
                from PIL import Image as _BPIL
                _bi = _BPIL.open(_bp)
                _bw, _bh = _bi.size
                _bd = _bi.info.get('dpi', (96, 96))
                _bdx = float(_bd[0]) if isinstance(_bd, tuple) else float(_bd)
                if _bdx < 1: _bdx = 96.0
                _bh_in = _bh / _bdx
                _section.different_first_page_header_footer = True
                _section.header_distance = Inches(0)
                _section.top_margin = Inches(_bh_in + 0.04)
                _section.left_margin  = Inches(0.5)
                _section.right_margin = Inches(0.5)
                _fph = _section.first_page_header
                for _p in _fph.paragraphs: _p.clear()
                _hp = _fph.paragraphs[0]
                _hp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                _pPr = _hp._p.get_or_add_pPr()
                _sp = OxmlElement('w:spacing'); _sp.set(qn('w:before'),'0'); _sp.set(qn('w:after'),'220'); _pPr.append(_sp)  # 220 twips ≈ 3.9mm
                _ind = OxmlElement('w:ind'); _ind.set(qn('w:left'),str(-720)); _ind.set(qn('w:right'),'0'); _ind.set(qn('w:hanging'),'0'); _pPr.append(_ind)
                _hr = _hp.add_run()
                _hr.add_picture(_bp, width=Inches(8.27))
                _nh = _section.header; _nh.is_linked_to_previous = False
                for _p in _nh.paragraphs: _p.clear()
            except Exception as _be:
                print(f"Banner Fehler: {_be}")
    
    # Seitenformat
    section = doc.sections[0]
    section.page_height = Inches(11.69)
    section.page_width = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Header-Variablen
    show_title     = starterlist.get("showTitle", "") if show_title_opt else ""
    comp_info      = starterlist.get("informationText", "")
    comp_number    = starterlist.get("competitionNumber", "")
    comp_title     = starterlist.get("competitionTitle", "")
    comp_subtitle  = starterlist.get("subtitle", "")
    division_number= starterlist.get("divisionNumber")

    # LOGO OBEN RECHTS - HINTER DEM TEXT
    if show_header:
        logo_path = starterlist.get("logoPath")
        logo_max_width_cm = starterlist.get("logoMaxWidthCm", 5.0)
        logo_inches = logo_max_width_cm / 2.54
        page_twips = int(7.27 * 1440)
        logo_twips = int((logo_inches + 0.1) * 1440)
        text_twips = page_twips - logo_twips

        if logo_path and os.path.exists(logo_path):
            h_table = doc.add_table(rows=1, cols=2)
            h_tbl = h_table._tbl
            h_tbl_pr = h_tbl.find(qn("w:tblPr"))
            if h_tbl_pr is None:
                h_tbl_pr = OxmlElement("w:tblPr"); h_tbl.insert(0, h_tbl_pr)
            h_borders = OxmlElement("w:tblBorders")
            for bn in ["top","left","bottom","right","insideH","insideV"]:
                b = OxmlElement(f"w:{bn}"); b.set(qn("w:val"), "none"); h_borders.append(b)
            h_tbl_pr.append(h_borders)
            h_tbl_w = OxmlElement("w:tblW")
            h_tbl_w.set(qn("w:w"), str(page_twips)); h_tbl_w.set(qn("w:type"), "dxa"); h_tbl_pr.append(h_tbl_w)
            tbl_grid = OxmlElement("w:tblGrid")
            for tw in [text_twips, logo_twips]:
                gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(tw)); tbl_grid.append(gc)
            h_tbl.insert(list(h_tbl).index(h_tbl_pr) + 1, tbl_grid)
            left_cell = h_table.rows[0].cells[0]
            right_cell = h_table.rows[0].cells[1]
            for cell, tw in zip([left_cell, right_cell], [text_twips, logo_twips]):
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = OxmlElement("w:tcW"); tc_w.set(qn("w:w"), str(tw)); tc_w.set(qn("w:type"), "dxa"); tc_pr.append(tc_w)
            first_tracker = [True]
            def _add_h(cell, text, sz, bold, sa=4):
                if first_tracker[0]:
                    p = cell.paragraphs[0]; p.clear(); first_tracker[0] = False
                else:
                    p = cell.add_paragraph()
                r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(sa)
                return p
            if show_title: _add_h(left_cell, show_title, 14, True, sa=2)
            if comp_number or comp_title:
                _cl = f"Competition {comp_number}"
                if comp_title: _cl += f" — {comp_title}"
                if division_number:
                    try:
                        _dn = int(division_number)
                        if _dn > 0: _cl += f" - {_dn}. Abt."
                    except: pass
                _add_h(left_cell, _cl, 11, True, sa=4)
            if comp_info: _add_h(left_cell, comp_info.lstrip("\n"), 10, False, sa=4)
            if comp_subtitle: _add_h(left_cell, comp_subtitle, 10, False, sa=4)
            # Logo rechts
            r_para = right_cell.paragraphs[0]
            r_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            r_para.paragraph_format.space_before = Pt(0)
            r_para.paragraph_format.space_after = Pt(0)
            r_run = r_para.add_run()
            r_run.add_picture(logo_path, width=Inches(logo_inches))
            # behindDoc=1
            try:
                import copy as _copy
                _wp = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
                _am = "http://schemas.openxmlformats.org/drawingml/2006/main"
                inline = r_run._r.find(f"{{{_wp}}}inline")
                if inline is not None:
                    anchor = OxmlElement("wp:anchor")
                    for k,v in [("distT","0"),("distB","0"),("distL","114300"),("distR","114300"),
                                 ("simplePos","0"),("relativeHeight","251658240"),("behindDoc","1"),
                                 ("locked","0"),("layoutInCell","1"),("allowOverlap","0")]:
                        anchor.set(k, v)
                    sp = OxmlElement("wp:simplePos"); sp.set("x","0"); sp.set("y","0"); anchor.append(sp)
                    posH = OxmlElement("wp:positionH"); posH.set("relativeFrom","margin")
                    alignH = OxmlElement("wp:align"); alignH.text = "right"; posH.append(alignH); anchor.append(posH)
                    posV = OxmlElement("wp:positionV"); posV.set("relativeFrom","margin")
                    alignV = OxmlElement("wp:align"); alignV.text = "top"; posV.append(alignV); anchor.append(posV)
                    extent = inline.find(f"{{{_wp}}}extent")
                    if extent is not None: anchor.append(_copy.deepcopy(extent))
                    eff = OxmlElement("wp:effectExtent")
                    for a in ["l","t","r","b"]: eff.set(a,"0")
                    anchor.append(eff)
                    anchor.append(OxmlElement("wp:wrapNone"))
                    docPr = inline.find(f"{{{_wp}}}docPr")
                    if docPr is not None: anchor.append(_copy.deepcopy(docPr))
                    anchor.append(OxmlElement("wp:cNvGraphicFramePr"))
                    a_graphic = inline.find(f"{{{_am}}}graphic")
                    if a_graphic is None:
                        for child in inline:
                            if child.tag.endswith("}graphic"): a_graphic = child; break
                    if a_graphic is not None: anchor.append(_copy.deepcopy(a_graphic))
                    drawing = inline.getparent()
                    if drawing is not None: drawing.replace(inline, anchor)
            except Exception as _be:
                print(f"behindDoc Fehler: {_be}")
        else:
            def _add_d(text, sz, bold, sa=4):
                p = doc.add_paragraph()
                r = p.add_run(text); r.font.size = Pt(sz); r.font.bold = bold
                p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(sa)
            if show_title: _add_d(show_title, 14, True, sa=2)
            if comp_number or comp_title:
                _cl = f"Competition {comp_number}"
                if comp_title: _cl += f" — {comp_title}"
                if division_number:
                    try:
                        _dn = int(division_number)
                        if _dn > 0: _cl += f" - {_dn}. Abt."
                    except: pass
                _add_d(_cl, 11, True, sa=4)
            if comp_info: _add_d(comp_info.lstrip("\n"), 10, False, sa=4)
            if comp_subtitle: _add_d(comp_subtitle, 10, False, sa=4)

    # Datum/Zeit/Ort für Starterliste-Zeile
    location = starterlist.get("location", "")
    start_time = starterlist.get("start", "")
    divisions = starterlist.get('divisions', [])
    division_number_val = starterlist.get("divisionNumber")
    start_raw = None
    if divisions and division_number_val is not None:
        try:
            div_num = int(division_number_val)
            for div in divisions:
                if div.get("number") == div_num:
                    if div.get("start"): start_raw = div.get("start"); break
        except: pass
    elif divisions:
        if divisions[0].get("start"): start_raw = divisions[0].get("start")
    if start_raw: start_time = start_raw

    date_text_always = format_header_datetime(start_time) if start_time else ""
    if date_text_always and location:
        date_text_always += f" - {location}"
    
    # Jury (wie im PDF: kompakte Zeile mit Badges)
    # WICHTIG: Reihenfolge ist IMMER E H C M B (egal wie viele Richter)
    judges = safe_get(starterlist, "judges", [])
    if judges:
        p_jury = doc.add_paragraph()
        run = p_jury.add_run("Jury: ")
        run.font.size = Pt(10)
        run.font.bold = True
        
        # Position-Mapping
        pos_map = {
            0: "E", 1: "H", 2: "C", 3: "M", 4: "B", 5: "K", 
            6: "V", 7: "S", 8: "R", 9: "P", 10: "F", 11: "A"
        }
        
        # Erstelle Dictionary: Position -> Judge
        judges_by_position = {}
        for judge in judges:
            position = judge.get("position", "")
            name = judge.get("name", "")
            
            if isinstance(position, int):
                pos_label = pos_map.get(position, str(position))
            else:
                pos_label = str(position) if position else ""
            
            judges_by_position[pos_label] = name
        
        # FESTE REIHENFOLGE: E H C M B (überspringen wenn nicht vorhanden)
        fixed_order = ["E", "H", "C", "M", "B"]
        
        for i, pos in enumerate(fixed_order):
            if pos in judges_by_position:
                # Positionsbuchstabe mit dunkelgrauem Hintergrund und weißer Schrift
                run = p_jury.add_run(f" {pos} ")
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)  # Weiß
                
                # Dunkelgrauer Hintergrund (wie Header)
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), '404040')  # Dunkelgrau
                run._element.get_or_add_rPr().append(shading)
                
                # Richtername (normal, kein Hintergrund)
                run = p_jury.add_run(f" {judges_by_position[pos]}")
                run.font.size = Pt(9)
                
                # Abstand zum nächsten Richter
                if i < len(fixed_order) - 1:
                    p_jury.add_run("  ")
    
    # "Starting Order" links + Datum rechtsbündig in einer Zeile
    _PAGE_TWIPS = int(7.27 * 1440)
    p_title = doc.add_paragraph()
    _pPr_so = p_title._p.get_or_add_pPr()
    # Paragraph-Einzug auf 0
    _ind_so = OxmlElement('w:ind')
    _ind_so.set(qn('w:left'), '0')
    _ind_so.set(qn('w:right'), '0')
    _pPr_so.append(_ind_so)
    # Tab-Stop rechts
    _tabs = OxmlElement('w:tabs')
    _tab  = OxmlElement('w:tab')
    _tab.set(qn('w:val'), 'right')
    _tab.set(qn('w:pos'), str(_PAGE_TWIPS))
    _tabs.append(_tab)
    _pPr_so.append(_tabs)
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after  = Pt(2)
    p_title.paragraph_format.left_indent  = 0
    p_title.paragraph_format.right_indent = 0
    run_so = p_title.add_run("Starting Order")
    run_so.font.bold = True
    run_so.font.size = Pt(12)
    if date_text_always:
        p_title.add_run("	")
        run_date = p_title.add_run(date_text_always)
        run_date.font.bold = True
        run_date.font.size = Pt(10)
    
    # Haupttabelle
    table = doc.add_table(rows=0, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Header Row - DUNKELGRAU mit weißer Schrift
    hdr_row = table.add_row()
    hdr_cells = hdr_row.cells
    
    headers = ["#", "Time", "CNO", "Horse", "Athlete", "Nat."]
    alignments = [
        WD_PARAGRAPH_ALIGNMENT.CENTER,  # #
        WD_PARAGRAPH_ALIGNMENT.CENTER,  # Time
        WD_PARAGRAPH_ALIGNMENT.CENTER,  # CNO
        WD_PARAGRAPH_ALIGNMENT.LEFT,    # Horse - LINKSBÜNDIG!
        WD_PARAGRAPH_ALIGNMENT.LEFT,    # Athlete - LINKSBÜNDIG!
        WD_PARAGRAPH_ALIGNMENT.CENTER   # Nat. - ZENTRIERT!
    ]
    
    for i, header_text in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        run = p.add_run(header_text)
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.size = Pt(9)  # 9pt für Header
        run.font.color.rgb = RGBColor(255, 255, 255)  # Weiße Schrift!
        p.alignment = alignments[i]  # Verwende korrekte Ausrichtung
        optimize_paragraph_spacing(p)
    
    set_dark_header_background(hdr_row)
    set_row_height_auto(hdr_row)
    
    # Starter Rows
    starters = safe_get(starterlist, "starters", [])
    
    # BREAKS Array holen und Map erstellen (wie in stream_abstammung!)
    breaks = safe_get(starterlist, "breaks", [])
    breaks_map = {}
    for br in breaks:
        try:
            # WICHTIG: afterNumberInCompetition kann 0 sein!
            after_num = br.get("afterNumberInCompetition")
            if after_num is None:
                after_num = -1
            else:
                after_num = int(after_num)
            breaks_map[after_num] = br
            print(f"🔍 Break registriert nach Starter #{after_num}: {br}")
        except:
            continue
    
    print(f"\n=== PAUSE-DEBUG ===")
    print(f"Anzahl Starter: {len(starters)}")
    print(f"Anzahl Breaks: {len(breaks)}")
    print(f"Breaks Map: {breaks_map}")
    
    # ROW COUNTER für korrektes Alternieren (inkl. Breaks!)
    row_counter = 0
    
    # FLAGGEN-PFAD: App läuft von C:\Python, Flaggen sind in C:\Python\flags
    def find_flag_path(nat_code):
        r"""Suche Flagge - App läuft von C:\Python aus"""
        
        possible_paths = [
            # Hauptpfade (App läuft von C:\Python)
            f"flags/{nat_code}.png",              # C:\Python\flags\GER.png
            f"./flags/{nat_code}.png",            # Explizit C:\Python\flags\GER.png
            
            # Absolute Windows-Pfade (Fallback)
            f"C:/Python/flags/{nat_code}.png",
            f"C:\\Python\\flags\\{nat_code}.png",
            
            # Falls doch woanders
            f"../flags/{nat_code}.png",
            f"../../flags/{nat_code}.png",
        ]
        
        for path in possible_paths:
            try:
                if os.path.exists(path):
                    abs_path = os.path.abspath(path)
                    print(f"✅ Flagge gefunden: {path} -> {abs_path}")
                    return path
            except Exception as e:
                continue
        
        print(f"⚠️ Flagge nicht gefunden für: {nat_code}")
        print(f"   Aktuelles Verzeichnis: {os.getcwd()}")
        print(f"   Gesuchte Pfade:")
        for p in possible_paths[:3]:
            abs_test = os.path.abspath(p)
            exists = os.path.exists(p)
            print(f"   - {p} -> {abs_test} {'✅' if exists else '❌'}")
        return None
    
    # WICHTIG: Prüfe ob es eine Pause VOR dem ersten Starter gibt (afterNumberInCompetition=0)
    if 0 in breaks_map:
        print("🎯 PAUSE VOR ERSTEM STARTER GEFUNDEN!")
        break_info = breaks_map[0]
        data_row = table.add_row()
        cells = data_row.cells
        set_row_height_auto(data_row, cant_split=True)
        
        # Pause-Text formatieren (wie andere Pausen)
        total_seconds = int(break_info.get("totalSeconds", 0) or 0)
        info_text = break_info.get("informationText", "")
        
        if total_seconds > 0:
            if total_seconds >= 3600:
                hours = total_seconds // 3600
                mins = (total_seconds % 3600) // 60
                break_text = f"Pause ({hours} h {mins:02d} min)"
            else:
                mins = total_seconds // 60
                break_text = f"Pause ({mins} min)"
            
            if info_text:
                break_text = f"{break_text} - {info_text}"
        else:
            break_text = info_text or "Pause"
        
        # Set break text in first cell
        p_break = cells[0].paragraphs[0]
        p_break.clear()
        run = p_break.add_run(break_text)
        run.font.name = 'Arial'  # Wie andere Pausen!
        run.font.size = Pt(10)
        run.font.bold = True  # Bold, nicht italic!
        p_break.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Clear other cells and merge (6 columns)
        for i in range(1, 6):
            cells[i].paragraphs[0].clear()
        
        cell_0 = cells[0]
        for i in range(1, 6):
            cell_0.merge(cells[i])
        
        # Hintergrund: Pause vor erstem Starter ist WEISS (nicht grau!)
        # Keine Hintergrundfarbe setzen = weiß
        # (Zebra-Logik: row_counter=0 → weiß, dann row_counter=1 → grau)
        
        row_counter += 1  # Nach Pause ist row_counter=1 → nächster Starter wird grau
        print(f"✅ Pause vor erstem Starter eingefügt (weiß): '{break_text}'")
    
    for idx, starter in enumerate(starters):
        data_row = table.add_row()
        cells = data_row.cells
        
        set_row_height_auto(data_row, cant_split=True)
        
        # Check if withdrawn FIRST!
        is_withdrawn = starter.get("withdrawn", False)
        is_hors_concours = starter.get("horsConcours", False)  # AK - Außer Konkurrenz
        
        # ALTERNATING GRAY basierend auf row_counter
        if row_counter % 2 == 1:
            set_light_gray_background(data_row)
        
        row_counter += 1  # Erhöhe nach jeder Datenzeile
        
        # Column 0: # (durchgestrichen wenn withdrawn) + AK darunter wenn horsConcours
        start_number = safe_get(starter, "startNumber", "")
        p0 = cells[0].paragraphs[0]
        p0.clear()
        run = p0.add_run(str(start_number))
        run.font.name = 'Arial'
        run.font.size = Pt(8)  # 8pt
        run.font.bold = False  # NICHT fett
        if is_withdrawn:
            run.font.strike = True
            # KEINE Farbänderung mehr!
        
        # AK unter der Startnummer wenn horsConcours
        if is_hors_concours:
            p0.add_run("\n")
            ak_run = p0.add_run("AK")
            ak_run.font.name = 'Arial'
            ak_run.font.size = Pt(7)
            ak_run.font.bold = True
            # Wenn withdrawn, auch AK durchstreichen
            if is_withdrawn:
                ak_run.font.strike = True
        
        p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        optimize_paragraph_spacing(p0)
        
        # Column 1: Time MIT SEKUNDEN (durchgestrichen wenn withdrawn)
        start_time = safe_get(starter, "startTime", "")
        formatted_time = format_time(start_time)
        p1 = cells[1].paragraphs[0]
        p1.clear()
        run = p1.add_run(formatted_time)
        run.font.name = 'Arial'
        run.font.size = Pt(8)  # 8pt
        run.font.bold = False  # NICHT fett
        if is_withdrawn:
            run.font.strike = True
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        optimize_paragraph_spacing(p1)
        
        # horses Array
        horses = starter.get("horses", [])
        horse = horses[0] if horses else {}
        
        # Column 2: CNO (KEINE Umrandung mehr!)
        cno = str(horse.get("cno", "")) if horse else ""
        p2 = cells[2].paragraphs[0]
        p2.clear()
        if cno:
            run = p2.add_run(cno)
            run.font.name = 'Arial'
            run.font.size = Pt(8)  # 8pt
            run.font.bold = False  # NICHT fett
            if is_withdrawn:
                run.font.strike = True
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        optimize_paragraph_spacing(p2)
        
        # Entferne ALLE Borders von ALLEN Zellen dieser Zeile
        for cell in cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            existing_borders = tcPr.xpath('.//w:tcBorders')
            for border in existing_borders:
                tcPr.remove(border)
        
        # Column 3: Horse (LINKSBÜNDIG!) - NEUE POSITION!
        horse_name = horse.get("name", "") if horse else ""
        
        p3 = cells[3].paragraphs[0]
        p3.clear()
        p3.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # LINKSBÜNDIG!
        
        # Pferdename in FETT (wie PDF: Arial, 9pt)
        if horse_name:
            run = p3.add_run(horse_name)
            run.font.name = 'Arial'
            run.font.size = Pt(9)  # 9pt fett
            run.font.bold = True
            if is_withdrawn:
                run.font.strike = True
            
            # BREED hinter Namen mit " - " (NICHT FETT!)
            breed = horse.get("breed", "") if horse else ""
            if breed:
                run = p3.add_run(f" - {breed}")
                run.font.name = 'Arial'
                run.font.size = Pt(9)  # 9pt
                run.font.bold = False  # NICHT FETT!
                if is_withdrawn:
                    run.font.strike = True
        
        if horse:
            horse_details = []
            
            # Alter mit "/" danach
            breeding_season = horse.get("breedingSeason")
            if breeding_season:
                try:
                    age = 2025 - int(breeding_season)
                    horse_details.append(f"{age}Y /")
                except:
                    pass
            
            # FARBE zwischen Alter und Geschlecht
            color = horse.get("color", "")
            if color:
                horse_details.append(f"{color} /")
            
            # Geschlecht mit "/" danach
            sex = horse.get("sex", "")
            if sex:
                horse_details.append(f"{sex.upper()} /")
            
            # Studbook
            studbook = horse.get("studbook", "")
            if studbook:
                horse_details.append(studbook)
            
            # Sire und DamSire mit "x" dazwischen
            sire = horse.get("sire", "")
            dam_sire = horse.get("damSire", "")
            
            if sire and dam_sire:
                horse_details.append(f"/ {sire} x {dam_sire}")
            elif sire:
                horse_details.append(f"/ {sire}")
            elif dam_sire:
                horse_details.append(f"/ {dam_sire}")
            
            # FEI-Nummer am Ende
            fei_number = horse.get("feiNumber", "")
            if fei_number:
                horse_details.append(f"/ {fei_number}")
            
            if horse_details:
                p3.add_run("\n")
                detail_text = " ".join(horse_details)
                run = p3.add_run(detail_text)
                run.font.name = 'Arial'
                run.font.size = Pt(7)  # 7pt für Details
                if is_withdrawn:
                    run.font.strike = True
            
            # Owner
            owner = horse.get("owner", "")
            if owner:
                p3.add_run("\n")
                run = p3.add_run(f"Owner: {owner}")
                run.font.name = 'Arial'
                run.font.size = Pt(7)  # 7pt
                run.font.italic = True
                if is_withdrawn:
                    run.font.strike = True
            
            # Breeder (NEU!)
            breeder = horse.get("breeder", "")
            if breeder:
                p3.add_run("\n")
                run = p3.add_run(f"Breeder: {breeder}")
                run.font.name = 'Arial'
                run.font.size = Pt(7)  # 7pt
                run.font.italic = True
                if is_withdrawn:
                    run.font.strike = True
        
        optimize_paragraph_spacing(p3)
        
        # Column 4: Athlete (LINKSBÜNDIG!) - NEUE POSITION!
        athlete = starter.get("athlete", {})
        athlete_name = athlete.get("name", "") if athlete else ""
        
        p4 = cells[4].paragraphs[0]
        p4.clear()
        p4.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # LINKSBÜNDIG!
        
        club = athlete.get("club", "") if athlete else ""
        nationality = athlete.get("nation", "") if athlete else ""
        club_or_country = ""
        if nationality and nationality.upper() != "GER":
            if not club or club.strip() == "" or club.strip().upper() == "GASTLIZENZ GER":
                club_or_country = get_country_name_english(nationality)
            else:
                club_or_country = club
        elif club:
            club_or_country = club

        if athlete_name:
            run = p4.add_run(athlete_name)
            run.font.name = 'Arial'
            run.font.size = Pt(9)
            run.font.bold = True
            if is_withdrawn:
                run.font.strike = True
        if club_or_country:
            p4.add_run("\n")
            run_club = p4.add_run(club_or_country)
            run_club.font.name = 'Arial'
            run_club.font.size = Pt(7)
            run_club.font.bold = False
            if is_withdrawn:
                run_club.font.strike = True

        optimize_paragraph_spacing(p4)
        
        # Column 5: Nat. mit echten Flaggen-Icons - NEUE POSITION!
        nationality = athlete.get("nation", "") if athlete else ""
        nat_code = get_nationality_code(nationality) if nationality else ""
        
        p5 = cells[5].paragraphs[0]
        p5.clear()
        if nat_code:
            # Versuche Flaggen-Icon zu finden
            flag_path = find_flag_path(nat_code)
            
            if flag_path:
                # Flaggen-Icon einfügen
                try:
                    from io import BytesIO
                    
                    abs_flag_path = os.path.abspath(flag_path)
                    
                    with open(abs_flag_path, 'rb') as f:
                        image_bytes = BytesIO(f.read())
                    
                    run = p5.add_run()
                    run.add_picture(image_bytes, width=Inches(0.25))
                    p5.add_run("\n")
                    
                    size = len(image_bytes.getvalue())
                    print(f"✅ Flagge eingefügt: {abs_flag_path} ({size} bytes)")
                except Exception as e:
                    print(f"⚠️ Konnte Flagge nicht laden: {flag_path} - {e}")
                    import traceback
                    traceback.print_exc()
            else:
                print(f"⚠️ Flagge nicht gefunden für: {nat_code}")
            
            # Code darunter
            run = p5.add_run(nat_code)
            run.font.name = 'Arial'
            run.font.size = Pt(7)  # 7pt
            run.font.bold = False  # NICHT fett
            if is_withdrawn:
                run.font.strike = True
            p5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        optimize_paragraph_spacing(p5)
        
        # BREAK HANDLING - Nach diesem Starter prüfen ob Pause kommt!
        try:
            starter_num = int(safe_get(starter, "startNumber", 0))
            if starter_num in breaks_map:
                print(f"\n✅ PAUSE nach Starter #{starter_num}!")
                break_info = breaks_map[starter_num]
                
                break_row = table.add_row()
                break_cells = break_row.cells
                
                set_row_height_auto(break_row)
                
                # HINTERGRUND ALTERNIEREND basierend auf row_counter
                # row_counter ist jetzt nach der letzten Zeile erhöht worden
                # Also: wenn row_counter gerade → letzte Zeile war ungerade (grau) → Break weiß
                #       wenn row_counter ungerade → letzte Zeile war gerade (weiß) → Break grau
                break_is_gray = (row_counter % 2 == 1)
                row_counter += 1  # Erhöhe auch für Break-Zeile!
                
                # Format break text
                total_seconds = int(break_info.get("totalSeconds", 0) or 0)
                info_text = break_info.get("informationText", "")
                
                if total_seconds > 0:
                    if total_seconds >= 3600:
                        hours = total_seconds // 3600
                        mins = (total_seconds % 3600) // 60
                        break_text = f"Pause ({hours} h {mins:02d} min)"
                    else:
                        mins = total_seconds // 60
                        break_text = f"Pause ({mins} min)"
                    
                    if info_text:
                        break_text = f"{break_text} - {info_text}"
                else:
                    break_text = info_text or "Pause"
                
                print(f"   Break Text: {break_text}, Gray: {break_is_gray}")
                
                # Set break text in first cell
                p_break = break_cells[0].paragraphs[0]
                p_break.clear()
                run = p_break.add_run(break_text)
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                p_break.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                optimize_paragraph_spacing(p_break)
                
                # Clear other cells and merge (6 columns)
                for i in range(1, 6):
                    break_cells[i].paragraphs[0].clear()
                
                cell_0 = break_cells[0]
                for i in range(1, 6):
                    cell_0.merge(break_cells[i])
                
                # Hintergrund nur wenn grau
                if break_is_gray:
                    tcPr = cell_0._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'E8E8E8')
                    tcPr.append(shd)
                
        except (ValueError, TypeError):
            pass
    
    # Tabelle formatieren
    set_column_widths(table)
    set_horizontal_borders_only(table)  # NUR horizontale Borders!
    set_cell_margins_tight(table)
    
    # SPONSORENLEISTE IM FOOTER (wie in word_abstammung_logo)
    sponsor_logo_path = "logos/sponsorenleiste.png"
    if os.path.exists(sponsor_logo_path):
        try:
            # Footer für alle Sections setzen
            for section in doc.sections:
                footer = _section.footer
                # Bei Banner auch first_page_footer mit Sponsorenleiste
                if show_banner and _section.different_first_page_header_footer:
                    try:
                        _fpf = _section.first_page_footer
                        _fpfp = _fpf.paragraphs[0] if _fpf.paragraphs else _fpf.add_paragraph()
                        _fpfp.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        _fpfp.clear()
                        _fpfr = _fpfp.add_run()
                        _fpfr.add_picture(sponsor_logo_path, width=Inches(7.5))
                        _fpfp.paragraph_format.space_after = Pt(0)
                        _fpfp.paragraph_format.space_before = Pt(0)
                    except: pass

                
                # Footer-Paragraph zentriert
                footer_para = footer.paragraphs[0]
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                footer_para.clear()
                
                # Sponsorenleiste in Footer einfügen
                footer_run = footer_para.add_run()
                footer_run.add_picture(sponsor_logo_path, width=Inches(7.5))  # Volle Breite
                
                # Footer-Formatierung optimieren
                footer_para.paragraph_format.space_after = Pt(0)
                footer_para.paragraph_format.space_before = Pt(0)
            
            print(f"WORD_INT DEBUG: Sponsorenleiste im Footer hinzugefügt: {sponsor_logo_path}")
        except Exception as e:
            print(f"WORD_INT DEBUG: Sponsorenleiste Footer-Fehler: {e}")
    else:
        print(f"WORD_INT DEBUG: Sponsorenleiste nicht gefunden: {sponsor_logo_path}")
    
    # Save to output directory
    print(f"WORD_INT DEBUG: Saving document to {output_path}")
    doc.save(output_path)
    print(f"WORD_INT DEBUG: Document saved successfully")
    return output_path

def create_stream(starterlist, template, filename):
    """Alias für Kompatibilität"""
    return render(starterlist, filename)
