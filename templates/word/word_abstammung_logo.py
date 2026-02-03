# templates/word/word_abstammung_logo.py - OPTIMIERT mit Zeilenhöhen-Steuerung
import os
import json
import shutil
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.enum.table import WD_TABLE_ALIGNMENT

# mapping English -> German for sex
SEX_MAP = {
    "MARE": "Stute",
    "GELDING": "Wallach", 
    "STALLION": "Hengst",
    "M": "Wallach",
    "F": "Stute",
    "": ""
}

WEEKDAY_MAP = {
    "Monday": "Montag", "Tuesday": "Dienstag", "Wednesday": "Mittwoch",
    "Thursday": "Donnerstag", "Friday": "Freitag", "Saturday": "Samstag",
    "Sunday": "Sonntag"
}

FIXED_WORD_TEMPLATE = "word_details_modern_simple.docx"
OUTPUT_DIR = "Ausgabe"

def _ensure_output_dir():
    """Stellt sicher, dass das Ausgabe-Verzeichnis existiert"""
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

def _safe_get(d, key, default=""):
    if not d:
        return default
    return d.get(key, default) if isinstance(d, dict) else default

def _format_time(iso):
    """Formatiert ISO-Zeit zu HH:MM:SS"""
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", ""))
        return dt.strftime("%H:%M:%S")
    except Exception:
        # Fallback: erste 8 Zeichen für vollständige Zeit
        return str(iso)[:8] if len(str(iso)) >= 8 else str(iso)

def _format_header_datetime(iso):
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", ""))
        weekday = WEEKDAY_MAP.get(dt.strftime("%A"), dt.strftime("%A"))
        return f"{weekday}, {dt.strftime('%d.%m.%Y um %H:%M Uhr')}"
    except Exception:
        return str(iso)

def _process_information_text(text):
    """Verarbeitet informationText - entfernt führende \n und macht Text fett"""
    if not text:
        return ""
    text = text.lstrip('\n')
    return text

def _set_header_gray_background(row):
    """Set gray background for header row"""
    tr = row._tr
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D9D9D9')  # Light gray
        tcPr.append(shd)

def _set_table_borders(table):
    """Add borders to all table cells"""
    tbl = table._tbl
    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.tcPr
            if tcPr is None:
                tcPr = OxmlElement("w:tcPr")
                cell.append(tcPr)
            
            tcBorders = OxmlElement('w:tcBorders')
            
            for border_name in ['w:top', 'w:left', 'w:bottom', 'w:right']:
                border = OxmlElement(border_name)
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcBorders.append(border)
            
            tcPr.append(tcBorders)

def _set_row_keep_together(row):
    """Prevent table row from breaking across pages"""
    tr = row._tr
    trPr = tr.trPr
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    
    # Add cantSplit property to prevent row from breaking across pages
    cantSplit = trPr.find(qn('w:cantSplit'))
    if cantSplit is None:
        cantSplit = OxmlElement('w:cantSplit')
        cantSplit.set(qn('w:val'), 'true')
        trPr.append(cantSplit)
    else:
        cantSplit.set(qn('w:val'), 'true')

def _set_row_height_auto(row):
    """Set row height to auto-adjust to content - AUS DRE_3 ÜBERNOMMEN"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    
    # Remove any existing height settings
    trHeight = trPr.find(qn('w:trHeight'))
    if trHeight is not None:
        trPr.remove(trHeight)
    
    # Set height to auto
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), '0')  # 0 = auto height
    trHeight.set(qn('w:hRule'), 'auto')  # auto rule
    trPr.append(trHeight)

def _set_cell_margins_tight(table):
    """Set tight cell margins to reduce padding - AUS DRE_3 ÜBERNOMMEN"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Remove existing cell margins
    tblCellMar = tblPr.find(qn('w:tblCellMar'))
    if tblCellMar is not None:
        tblPr.remove(tblCellMar)
    
    # Set tight cell margins
    tblCellMar = OxmlElement('w:tblCellMar')
    
    # Top margin: 36 twips (about 2pt)
    top = OxmlElement('w:top')
    top.set(qn('w:w'), '36')
    top.set(qn('w:type'), 'dxa')
    tblCellMar.append(top)
    
    # Left margin: 72 twips (about 4pt)
    left = OxmlElement('w:left')
    left.set(qn('w:w'), '72')
    left.set(qn('w:type'), 'dxa')
    tblCellMar.append(left)
    
    # Bottom margin: 36 twips (about 2pt)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:w'), '36')
    bottom.set(qn('w:type'), 'dxa')
    tblCellMar.append(bottom)
    
    # Right margin: 72 twips (about 4pt)
    right = OxmlElement('w:right')
    right.set(qn('w:w'), '72')
    right.set(qn('w:type'), 'dxa')
    tblCellMar.append(right)
    
    tblPr.append(tblCellMar)

def _optimize_paragraph_spacing(paragraph):
    """Optimize paragraph spacing for compact display - AUS DRE_3 ÜBERNOMMEN"""
    pPr = paragraph._element.get_or_add_pPr()
    
    # Remove existing spacing
    spacing = pPr.find(qn('w:spacing'))
    if spacing is not None:
        pPr.remove(spacing)
    
    # Set tight spacing
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')     # No space before
    spacing.set(qn('w:after'), '0')      # No space after
    spacing.set(qn('w:line'), '240')     # Line spacing: 240 twips = 12pt (single)
    spacing.set(qn('w:lineRule'), 'auto') # Auto line rule
    pPr.append(spacing)

def _set_header_row_repeat(table):
    """Enable header row to repeat on each new page"""
    tbl = table._tbl
    
    # Get the first row (header row)
    if len(tbl.tr_lst) > 0:
        header_row = tbl.tr_lst[0]
        
        # Set the tblHeader property to repeat the header row on new pages
        trPr = header_row.trPr
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            header_row.insert(0, trPr)
        
        # Add or update tblHeader element
        tblHeader = trPr.find(qn('w:tblHeader'))
        if tblHeader is None:
            tblHeader = OxmlElement('w:tblHeader')
            tblHeader.set(qn('w:val'), 'true')
            trPr.append(tblHeader)
        else:
            tblHeader.set(qn('w:val'), 'true')
        
        print("WORD DEBUG: Header row repeat enabled")

def _set_judges_table_widths(table):
    """Set fixed column widths for judges table"""
    tbl = table._tbl
    
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblPr.remove(tblStyle)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    tblGrid = OxmlElement('w:tblGrid')
    judges_column_widths_twips = [360, 7200]
    
    for width in judges_column_widths_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)
    
    if tblPr.getparent() is not None:
        tblPr.getparent().insert(tblPr.getparent().index(tblPr) + 1, tblGrid)
    else:
        tbl.insert(0, tblGrid)
    
    for row in tbl.tr_lst:
        for i, cell in enumerate(row.tc_lst):
            if i < len(judges_column_widths_twips):
                tcPr = cell.tcPr
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell.insert(0, tcPr)
                
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(judges_column_widths_twips[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def _set_abstammung_column_widths(table):
    """Set table to use exact column widths for 5-column abstammung layout (Start|KoNr|Reiter|Pferd|Ergeb) - ANGEPASST"""
    tbl = table._tbl
    
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblStyle = tblPr.find(qn('w:tblStyle'))
    if tblStyle is not None:
        tblPr.remove(tblStyle)
    
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is not None:
        tblPr.remove(tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '0')
    tblW.set(qn('w:type'), 'auto')
    tblPr.append(tblW)
    
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        tbl.remove(tblGrid)
    
    tblGrid = OxmlElement('w:tblGrid')
    # FINAL optimierte 5-column widths: Start (etwas breiter) | KoNr | Reiter | Pferd | Ergeb
    abstammung_column_widths_twips = [900, 540, 3150, 5750, 800]  # Start +100, Reiter -50, Pferd -50
    
    for width in abstammung_column_widths_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(width))
        tblGrid.append(gridCol)
    
    if tblPr.getparent() is not None:
        tblPr.getparent().insert(tblPr.getparent().index(tblPr) + 1, tblGrid)
    else:
        tbl.insert(0, tblGrid)
    
    for row in tbl.tr_lst:
        for i, cell in enumerate(row.tc_lst):
            if i < len(abstammung_column_widths_twips):
                tcPr = cell.tcPr
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell.insert(0, tcPr)
                
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is not None:
                    tcPr.remove(tcW)
                
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(abstammung_column_widths_twips[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)

def _get_judge_data_for_display_abstammung(judges):
    """
    Prozessiert Richter für die Anzeige - alle Positionen für die Richterliste in korrekter Reihenfolge
    """
    pos_map = {
        0: "E", 1: "H", 2: "C", 3: "M", 4: "B", 5: "K", 6: "V", 
        7: "S", 8: "R", 9: "P", 10: "F", 11: "A",
        "WARM_UP_AREA": "Aufsicht", "WATER_JUMP": "Wasser"
    }
    
    dressage_positions = ["E", "H", "C", "M", "B"]
    judges_by_position = {}  # Dictionary of Lists!
    other_judges = []
    
    for judge in judges:
        position = judge.get("position", "")
        if isinstance(position, int):
            pos_label = pos_map.get(position, str(position))
        else:
            pos_label = pos_map.get(str(position), str(position))
        
        judge_copy = judge.copy()
        judge_copy["pos_label"] = pos_label
        
        if pos_label in dressage_positions:
            # Füge zu Liste hinzu statt zu überschreiben
            if pos_label not in judges_by_position:
                judges_by_position[pos_label] = []
            judges_by_position[pos_label].append(judge_copy)
        else:
            other_judges.append(judge_copy)
    
    # Return judges in E H C M B order first, then others
    result = []
    for pos in dressage_positions:
        if pos in judges_by_position:
            # Füge ALLE Richter dieser Position hinzu
            result.extend(judges_by_position[pos])
    
    # Add other judges sorted by position label
    other_judges.sort(key=lambda j: j["pos_label"])
    result.extend(other_judges)
    
    return result
    
    
def render(starterlist: dict, filename: str):
    """
    Creates a Word document with 5-column abstammung layout (Start|KoNr|Reiter|Pferd|Ergeb) with prominent logo
    """
    print("WORD ABSTAMMUNG LOGO DEBUG: Starting render function")
    
    if starterlist is None:
        raise ValueError("starterlist is None")

    if isinstance(starterlist, str):
        try:
            starterlist = json.loads(starterlist)
        except Exception as e:
            raise ValueError(f"starterlist is a string but not valid JSON: {e}")

    _ensure_output_dir()
    
    # Use the fixed Word template as base
    template_path = os.path.join("templates", "word", FIXED_WORD_TEMPLATE)
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Fixed Word template not found: {template_path}")
    
    # Output path
    output_path = os.path.join(OUTPUT_DIR, filename)
    
    # Copy template to output
    shutil.copy2(template_path, output_path)
    
    # Open the copied template document
    print("WORD ABSTAMMUNG LOGO DEBUG: Loading document from template")
    doc = Document(output_path)
    
    # SEITENRÄNDER ANPASSEN - SEHR schmal für maximalen Platz
    for section in doc.sections:
        section.left_margin = Inches(0.2)   # Noch schmaler - nur 0.2 Zoll
        section.right_margin = Inches(0.2)  # Noch schmaler - nur 0.2 Zoll
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
    
    # Find and replace STARTER_TABLE placeholder
    print("WORD DEBUG: Looking for STARTER_TABLE placeholder")
    starter_table_found = False
    placeholder_paragraph = None
    
    for paragraph in doc.paragraphs:
        if "STARTER_TABLE" in paragraph.text:
            print("WORD DEBUG: Found STARTER_TABLE placeholder")
            placeholder_paragraph = paragraph
            starter_table_found = True
            break
    
    if not starter_table_found:
        print("WORD DEBUG: STARTER_TABLE placeholder not found, creating at end of document")
        placeholder_paragraph = doc.add_paragraph("STARTER_TABLE")

    # Extract data from starterlist for header generation
    show_title = starterlist.get("showTitle", "")
    comp_title = starterlist.get("competitionTitle", "")
    comp_number = starterlist.get("competitionNumber", "")
    comp_subtitle = starterlist.get("subtitle", "")
    comp_info = starterlist.get("informationText", "")
    location = starterlist.get("location", "")
    start_time = starterlist.get("start", "")
    division_number = starterlist.get("divisionNumber")

    # Handle division start times - 3-Szenario Logik
    start_raw = None
    divisions = starterlist.get('divisions', [])
    
    if divisions and division_number is not None:
        try:
            div_num = int(division_number)
            for div in divisions:
                if div.get("number") == div_num:
                    division_start = div.get("start")
                    if division_start:
                        start_raw = division_start
                        break
        except (ValueError, TypeError):
            pass
    elif divisions and len(divisions) > 0:
        first_division = divisions[0]
        division_start = first_division.get("start")
        if division_start:
            start_raw = division_start
    
    if not start_raw:
        start_raw = start_time
    
    if start_raw:
        start_time = start_raw

    # Remove the placeholder paragraph first
    placeholder_paragraph._element.getparent().remove(placeholder_paragraph._element)

    # LOGO DIREKT AM ANFANG - Prüfungsspezifisches Logo-System
    logo_path = starterlist.get("logoPath")
    if logo_path and os.path.exists(logo_path):
        try:
            # Logo-Paragraph ganz oben, rechtsbündig
            logo_para = doc.add_paragraph()
            logo_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            logo_run = logo_para.add_run()
            logo_run.add_picture(logo_path, width=Inches(1.2))  # Größeres Logo
            logo_para.paragraph_format.space_after = Pt(0)
            logo_para.paragraph_format.space_before = Pt(0)
            
            # Negativer Spacer um das Logo nach oben zu "ziehen"
            spacer_para = doc.add_paragraph()
            spacer_para.paragraph_format.space_before = Pt(0)  # Kein negativer Wert
            spacer_para.paragraph_format.space_after = Pt(0)
            
            print(f"WORD ABSTAMMUNG LOGO DEBUG: Logo positioned at top right: {logo_path}")
        except Exception as e:
            print(f"WORD ABSTAMMUNG LOGO DEBUG: Logo error: {e}")

    # KOMPLETTER HEADER ABER KOMPAKT - MIT LOGO
    print("WORD DEBUG: Creating COMPLETE but COMPACT header")

    if show_title:
        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(show_title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title_para.paragraph_format.space_after = Pt(0)
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.line_spacing = 1.0

    # Add competition info - COMPACT
    if comp_info:
        processed_info_text = _process_information_text(comp_info)
        info_para = doc.add_paragraph()
        run = info_para.add_run(processed_info_text)
        run.font.size = Pt(14)
        run.font.bold = True
        info_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        info_para.paragraph_format.space_before = Pt(0)
        info_para.paragraph_format.space_after = Pt(0)
        info_para.paragraph_format.line_spacing = 1.0

    # Add competition line - COMPACT mit korrekter Encoding
    if comp_number or comp_title:
        comp_line = f"Prüfung {comp_number}"
        if comp_title:
            comp_line += f" — {comp_title}"  # Korrekte deutsche Umlaute
        if division_number:
            try:
                div_num = int(division_number)
                if div_num > 0:
                    comp_line += f" - {div_num}. Abt." if "abt" not in comp_title.lower() else ""
            except:
                pass
        
        comp_para = doc.add_paragraph()
        run = comp_para.add_run(comp_line)
        run.font.size = Pt(12)
        run.font.bold = True
        comp_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        comp_para.paragraph_format.space_before = Pt(0)
        comp_para.paragraph_format.space_after = Pt(0)
        comp_para.paragraph_format.line_spacing = 1.0

    # Add subtitle - COMPACT
    if comp_subtitle:
        sub_para = doc.add_paragraph()
        run = sub_para.add_run(comp_subtitle)
        run.font.size = Pt(11)
        sub_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        sub_para.paragraph_format.space_before = Pt(0)
        sub_para.paragraph_format.space_after = Pt(0)
        sub_para.paragraph_format.line_spacing = 1.0

    # Add date/location - MINIMAL spacing to table
    if start_time:
        date_text = _format_header_datetime(start_time)
        if location:
            date_text += f" - {location}"
        date_para = doc.add_paragraph()
        run = date_para.add_run(date_text)
        run.font.size = Pt(11)
        date_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(3)  # MINIMAL spacing to table
        date_para.paragraph_format.line_spacing = 1.0

    # Create table DIREKT nach kompaktem Header - 5 COLUMNS for abstammung layout (Start|KoNr|Reiter|Pferd|Ergeb)
    print("WORD DEBUG: Creating table after compact header - 5 columns for abstammung")
    table = doc.add_table(rows=0, cols=5)  # 5 columns: Start, KoNr, Reiter, Pferd/Abstammung, Ergeb.
    
    # Try to set table style, but continue without if not available
    try:
        table.style = 'Table Grid'
        print("WORD DEBUG: Applied Table Grid style")
    except Exception as e:
        print(f"WORD DEBUG: Table Grid style not available: {e}, continuing without style")
    
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    # Add header row
    hdr_row = table.add_row()
    hdr_cells = hdr_row.cells
    
    headers = ["Start", "KoNr.", "Reiter", "Pferd", "Ergeb."]
    
    for i, header_text in enumerate(headers):
        if i >= len(hdr_cells):
            break
            
        p = hdr_cells[i].paragraphs[0]
        p.clear()
        
        if i == 3:  # Pferd column (4th column)
            run = p.add_run("Pferd")
            run.font.bold = True
            run.font.size = Pt(9)
            
            p.add_run("\n")
            run = p.add_run("Geboren / Farbe / Geschlecht / Vater x Muttervater")
            run.font.bold = False
            run.font.size = Pt(8)
            
            p.add_run("\n")
            run = p.add_run("Besitzer / Züchter")
            run.font.bold = False
            run.font.size = Pt(8)
            
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            _optimize_paragraph_spacing(p)  # AUS DRE_3 ÜBERNOMMEN
        else:
            run = p.add_run(str(header_text))
            run.font.bold = True
            run.font.size = Pt(9)
            if i == 2:  # Reiter column - linksbündig
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _optimize_paragraph_spacing(p)  # AUS DRE_3 ÜBERNOMMEN
    
    # Apply gray background to header row
    _set_header_gray_background(hdr_row)

    _set_abstammung_column_widths(table)
    
    # Enable header row repeat on new pages
    _set_header_row_repeat(table)

    # Process starters and breaks mit Abteilungslogik
    starters = starterlist.get("starters", [])
    breaks = starterlist.get("breaks", [])
    
    print(f"WORD DEBUG: Found {len(starters)} starters in starterlist")
    print(f"WORD DEBUG: Found {len(breaks)} breaks in starterlist")
    
    breaks_map = {}
    for br in breaks:
        try:
            after_num = int(br.get("afterNumberInCompetition", -1))
            breaks_map[after_num] = br
        except:
            continue

    # Gruppierung nach groupNumber für Abteilungen
    starters_by_group = {}
    for starter in starters:
        group_num = starter.get("groupNumber", 0)
        if group_num not in starters_by_group:
            starters_by_group[group_num] = []
        starters_by_group[group_num].append(starter)
    
    # Gruppen sortiert verarbeiten
    for group_num in sorted(starters_by_group.keys()):
        group_starters = starters_by_group[group_num]
        
        # Abteilungsheader hinzufügen (wenn mehr als eine Gruppe)
        if len(starters_by_group) > 1 and group_num > 0:
            header_row = table.add_row()
            header_cells = header_row.cells
            
            # Optimize division header row height - AUS DRE_3 ÜBERNOMMEN
            _set_row_height_auto(header_row)
            
            # Abteilungstext linksbündig
            p_header = header_cells[0].paragraphs[0]
            p_header.clear()
            run = p_header.add_run(f"{group_num}. Abteilung")
            run.font.size = Pt(11)
            run.font.bold = True
            p_header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Linksbündig
            _optimize_paragraph_spacing(p_header)  # AUS DRE_3 ÜBERNOMMEN
            
            # Merge über alle Spalten
            cell_0 = header_cells[0]
            for i in range(1, 5):
                header_cells[i].paragraphs[0].clear()
                cell_0.merge(header_cells[i])
            
            # Grauer Hintergrund für Abteilungsheader
            tcPr = cell_0._tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'E6E6E6')  # Hellgrau
            tcPr.append(shd)
        
        # Erste Startzeit pro Gruppe bestimmen
        group_start_time = None
        for starter in group_starters:
            start_time_candidate = starter.get("startTime")
            if start_time_candidate:
                group_start_time = start_time_candidate
                break

        # Starter der Gruppe verarbeiten
        for starter_idx, starter in enumerate(group_starters):
            row = table.add_row()
            row_cells = row.cells
            
            # Prevent this row from breaking across pages
            _set_row_keep_together(row)
            # NEUE FUNKTIONIERENDE Zeilenhöhen-Optimierung - AUS DRE_3 ÜBERNOMMEN
            _set_row_height_auto(row)
            
            # Column 0: Start number and time
            start_num = starter.get("startNumber", "")
            
            p0 = row_cells[0].paragraphs[0]
            p0.clear()
            p0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _optimize_paragraph_spacing(p0)  # AUS DRE_3 ÜBERNOMMEN
            
            if start_num:
                run = p0.add_run(str(start_num))
                run.font.bold = True
                run.font.size = Pt(9)
            
            # Zeit anzeigen: Entweder bei jedem Starter (wenn keine Abteilungen) oder nur beim ersten der Gruppe
            if len(starters_by_group) == 1:  # Keine Abteilungen - zeige immer die Zeit
                individual_start_time = starter.get("startTime")
                if individual_start_time:
                    start_time_str = _format_time(individual_start_time)
                    if start_time_str:
                        p0.add_run("\n")
                        run = p0.add_run(start_time_str)
                        run.font.size = Pt(9)  # Zeit-Schriftgröße von 8 auf 9
            else:  # Abteilungen vorhanden - nur erste Zeit pro Gruppe
                if starter_idx == 0 and group_start_time:
                    start_time_str = _format_time(group_start_time)
                    if start_time_str:
                        p0.add_run("\n")
                        run = p0.add_run(start_time_str)
                        run.font.size = Pt(9)  # Zeit-Schriftgröße von 8 auf 9

            # Column 1: KoNr.
            horses = starter.get("horses", [])
            p1 = row_cells[1].paragraphs[0]
            p1.clear()
            p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            _optimize_paragraph_spacing(p1)  # AUS DRE_3 ÜBERNOMMEN
            
            if horses:
                horse = horses[0]
                cno = _safe_get(horse, "cno", "")
                if cno:
                    run = p1.add_run(str(cno))
                    run.font.size = Pt(9)

            # Column 2: Reiter + Verein + Nation (erweitert)
            athlete = starter.get("athlete", {})
            rider_name = _safe_get(athlete, "name", "")
            club = _safe_get(athlete, "club", "")
            nation = _safe_get(athlete, "nation", "")
            
            p2 = row_cells[2].paragraphs[0]
            p2.clear()
            p2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            _optimize_paragraph_spacing(p2)  # AUS DRE_3 ÜBERNOMMEN
            
            # Reitername (KEINE Großschreibung)
            if rider_name:
                run = p2.add_run(rider_name)  # Entfernt .upper()
                run.font.bold = True
                run.font.size = Pt(9)
            
            # Verein
            if club:
                p2.add_run("\n")
                run = p2.add_run(club)
                run.font.size = Pt(8)
            
            # Nation
            if nation:
                p2.add_run("\n")
                run = p2.add_run(nation)
                run.font.size = Pt(8)

            # Column 3: Pferd + Abstammung info (ERWEITERT für Abstammung)
            p3 = row_cells[3].paragraphs[0]
            p3.clear()
            _optimize_paragraph_spacing(p3)  # AUS DRE_3 ÜBERNOMMEN
            
            content_parts = []
            
            if horses:
                horse = horses[0]
                horse_name = _safe_get(horse, "name", "")
                studbook = _safe_get(horse, "studbook", "")
                
                horse_line = ""
                if horse_name:
                    horse_line = horse_name  # Entfernt .upper()
                if studbook:
                    if horse_line:
                        horse_line += f" - {studbook}"
                    else:
                        horse_line = studbook
                if horse_line:
                    content_parts.append(horse_line)
                
                # ERWEITERTE Details für Abstammung - KORREKTE REIHENFOLGE
                details_parts = []
                year = _safe_get(horse, "breedingSeason", "")
                if year:
                    details_parts.append(str(year))
                
                color = _safe_get(horse, "color", "")
                if color:
                    details_parts.append(color)
                
                sex = _safe_get(horse, "sex", "")
                sex_german = SEX_MAP.get(str(sex).upper(), sex)
                if sex_german:
                    details_parts.append(sex_german)
                    
                sire = _safe_get(horse, "sire", "")
                damsire = _safe_get(horse, "damSire", "")
                if sire:
                    pedigree = sire
                    if damsire:
                        pedigree += f" x {damsire}"
                    details_parts.append(pedigree)
                
                if details_parts:
                    details_text = " / ".join(details_parts)
                    content_parts.append(details_text)
                
                # Besitzer / Züchter - WIEDERHERGESTELLT!
                owner = _safe_get(horse, "owner", "")
                breeder = _safe_get(horse, "breeder", "")
                
                owner_breeder_parts = []
                if owner:
                    owner_breeder_parts.append(f"B: {owner}")
                if breeder:
                    owner_breeder_parts.append(f"Z: {breeder}")
                
                if owner_breeder_parts:
                    owner_breeder_text = " / ".join(owner_breeder_parts)
                    content_parts.append(owner_breeder_text)
                
                # ZUSÄTZLICHE Abstammungsdetails falls vorhanden
                # Dam (Mutter)
                dam = _safe_get(horse, "dam", "")
                if dam:
                    content_parts.append(f"Mutter: {dam}")
            
            for i, part in enumerate(content_parts):
                if i > 0:
                    p3.add_run("\n")
                
                if i == 0:  # Pferdename + Studbook
                    # Split horse line to handle name and studbook separately
                    if " - " in part:
                        horse_name_part, studbook_part = part.split(" - ", 1)
                        # Horse name - fett
                        run = p3.add_run(horse_name_part)
                        run.font.bold = True
                        run.font.size = Pt(9)
                        # Studbook - nicht fett
                        run = p3.add_run(f" - {studbook_part}")
                        run.font.bold = False
                        run.font.size = Pt(9)
                    else:
                        # Nur Pferdename, fett
                        run = p3.add_run(part)
                        run.font.bold = True
                        run.font.size = Pt(9)
                else:  # Details
                    run = p3.add_run(part)
                    run.font.size = Pt(8)

            # Column 4: Ergeb. (empty for results)
            p4 = row_cells[4].paragraphs[0]
            p4.clear()
            p4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Handle withdrawn entries
            if starter.get("withdrawn", False):
                for cell in row_cells:
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'D9D9D9')
                    tcPr.append(shd)
                    
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(96, 96, 96)
                            run.font.strike = True

            # Check for breaks
            try:
                starter_num = int(start_num)
                if starter_num in breaks_map:
                    break_info = breaks_map[starter_num]
                    
                    break_row = table.add_row()
                    break_cells = break_row.cells
                    
                    # Prevent break row from splitting across pages
                    _set_row_keep_together(break_row)
                    
                    total_seconds = int(break_info.get("totalSeconds", 0) or 0)
                    info_text = break_info.get("informationText", "")
                    
                    if total_seconds > 0:
                        if total_seconds >= 3600:
                            hours = total_seconds // 3600
                            mins = (total_seconds % 3600) // 60
                            time_text = f"{hours} Std. {mins:02d} Min. Pause"
                        else:
                            mins = total_seconds // 60
                            time_text = f"{mins} Minuten Pause"
                        
                        break_text = f"{time_text} — {info_text}" if info_text else time_text
                    else:
                        break_text = info_text or "Pause"
                    
                    p_break = break_cells[0].paragraphs[0]
                    p_break.clear()
                    run = p_break.add_run(break_text)
                    run.font.size = Pt(10)
                    run.font.bold = True
                    p_break.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    
                    # Clear other cells and merge (5 columns for Abstammung)
                    for i in range(1, 5):
                        break_cells[i].paragraphs[0].clear()
                    
                    cell_0 = break_cells[0]
                    for i in range(1, 5):
                        cell_0.merge(break_cells[i])
                    
                    tcPr = cell_0._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), 'D9D9D9')
                    tcPr.append(shd)
                    
            except (ValueError, TypeError):
                pass

    _set_table_borders(table)

    # Add judges section - WICHTIG: Mit korrekter Reihenfolge E H C M B
    judges = starterlist.get("judges", [])
    if judges:
        doc.add_paragraph()
        
        p = doc.add_paragraph()
        run = p.add_run("Richter:")
        run.font.bold = True
        run.font.size = Pt(11)
        
        # Get judge data for display - KORREKTE REIHENFOLGE E H C M B
        judge_display_data = _get_judge_data_for_display_abstammung(judges)
        
        if judge_display_data:
            judges_table = doc.add_table(rows=1 + len(judge_display_data), cols=2)
            
            # Try to set table style, but continue without if not available
            try:
                judges_table.style = 'Table Grid'
                print("WORD DEBUG: Applied Table Grid style to judges table")
            except Exception as e:
                print(f"WORD DEBUG: Table Grid style not available for judges table: {e}, continuing without style")
            
            judges_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            # Header row
            hdr_cells = judges_table.rows[0].cells
            headers = ["Pos", "Richter"]
            
            for i, header_text in enumerate(headers):
                p = hdr_cells[i].paragraphs[0]
                p.clear()
                run = p.add_run(header_text)
                run.font.bold = True
                run.font.size = Pt(9)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Apply gray background to judges header row
            _set_header_gray_background(judges_table.rows[0])
            
            # Judge rows - KORREKTE REIHENFOLGE E H C M B
            for i, judge_data in enumerate(judge_display_data):
                row = judges_table.rows[i + 1]
                cells = row.cells
                
                # Position
                p = cells[0].paragraphs[0]
                p.clear()
                run = p.add_run(judge_data.get("pos_label", ""))
                run.font.size = Pt(9)
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                
                # Name
                p = cells[1].paragraphs[0]
                p.clear()
                run = p.add_run(judge_data.get("name", ""))
                run.font.size = Pt(9)

            _set_table_borders(judges_table)
            _set_judges_table_widths(judges_table)

    # SPONSORENLEISTE IM FUß JEDER SEITE - KORREKT AM ENDE
    sponsor_logo_path = os.path.join("logos", "sponsorenleiste.png")
    if os.path.exists(sponsor_logo_path):
        try:
            # Footer für alle Sections setzen
            for section in doc.sections:
                footer = section.footer
                
                # Footer-Paragraph zentriert
                footer_para = footer.paragraphs[0]
                footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                footer_para.clear()
                
                # Sponsorenleiste in Footer einfügen
                footer_run = footer_para.add_run()
                footer_run.add_picture(sponsor_logo_path, width=Inches(7.5))  # Volle verfügbare Breite
                
                # Footer-Formatierung optimieren
                footer_para.paragraph_format.space_after = Pt(0)
                footer_para.paragraph_format.space_before = Pt(0)
            
            print(f"WORD ABSTAMMUNG LOGO DEBUG: Sponsorenleiste im Footer hinzugefügt: {sponsor_logo_path}")
        except Exception as e:
            print(f"WORD ABSTAMMUNG LOGO DEBUG: Sponsorenleiste Footer-Fehler: {e}")
    else:
        print(f"WORD ABSTAMMUNG LOGO DEBUG: Sponsorenleiste nicht gefunden: {sponsor_logo_path}")

    # Save document
    print("WORD DEBUG: Starting document save")
    try:
        doc.save(output_path)
        print(f"WORD ABSTAMMUNG LOGO DEBUG: Document saved to {output_path}")
        return output_path
    except Exception as e:
        print(f"WORD DEBUG: Error saving document: {e}")
        alt_output_path = filename
        try:
            doc.save(alt_output_path)
            print(f"WORD ABSTAMMUNG LOGO DEBUG: Document saved to alternative location: {alt_output_path}")
            return alt_output_path
        except Exception as e2:
            print(f"WORD DEBUG: Failed to save to alternative location: {e2}")
            raise